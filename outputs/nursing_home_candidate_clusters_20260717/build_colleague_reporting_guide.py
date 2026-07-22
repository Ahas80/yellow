#!/usr/bin/env python3

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path("/Users/Aamir/Desktop/rUTIs/outputs/nursing_home_candidate_clusters_20260717")
OUTPUT_PATH = BASE_DIR / "Colleague_email_and_reporting_guide.docx"

NAVY = "16324F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8F1F8"
PALE_GREEN = "E2F0D9"
PALE_AMBER = "FFF2CC"
PALE_RED = "FCE4D6"
PALE_GRAY = "F2F4F7"
MID_GRAY = "5B6573"
LIGHT_BORDER = "D7DEE8"
WHITE = "FFFFFF"
BLACK = "1F2937"


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=LIGHT_BORDER, size=4):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_inches):
    width_twips = int(width_inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_keep_with_next(paragraph, keep=True):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = p_pr.find(qn("w:keepNext"))
    if keep and keep_next is None:
        keep_next = OxmlElement("w:keepNext")
        p_pr.append(keep_next)
    elif not keep and keep_next is not None:
        p_pr.remove(keep_next)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])


def add_labelled_paragraph(document, label, text, style=None):
    paragraph = document.add_paragraph(style=style)
    label_run = paragraph.add_run(label)
    label_run.bold = True
    label_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    paragraph.add_run(text)
    return paragraph


def add_callout(document, text, fill=PALE_BLUE, border=BLUE):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=border, size=8)
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.5)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, top=130, start=180, bottom=130, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    return table


def add_section_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    set_keep_with_next(paragraph)
    return paragraph


def format_date(value):
    return value if value else "Not available"


with (BASE_DIR / "analysis_summary.json").open() as handle:
    summary = json.load(handle)

with (BASE_DIR / "cluster_summary.csv").open(newline="") as handle:
    clusters = list(csv.DictReader(handle))

document = Document()
section = document.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = document.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

title_style = styles["Title"]
title_style.font.name = "Calibri"
title_style.font.size = Pt(23)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor.from_string(NAVY)
title_style.paragraph_format.space_after = Pt(5)

subtitle_style = styles["Subtitle"]
subtitle_style.font.name = "Calibri"
subtitle_style.font.size = Pt(12)
subtitle_style.font.color.rgb = RGBColor.from_string(MID_GRAY)
subtitle_style.paragraph_format.space_after = Pt(16)

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

# Memo-style running header.
header = section.header
header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
header_table.autofit = False
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(header_table, color=NAVY, size=0)
header_left, header_right = header_table.rows[0].cells
set_cell_width(header_left, 4.7)
set_cell_width(header_right, 1.8)
for cell in (header_left, header_right):
    set_cell_fill(cell, NAVY)
    set_cell_margins(cell, top=80, start=120, bottom=80, end=120)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
left_p = header_left.paragraphs[0]
left_p.paragraph_format.space_after = Pt(0)
left_run = left_p.add_run("GENOMIC CLUSTER REVIEW")
left_run.bold = True
left_run.font.size = Pt(9)
left_run.font.color.rgb = RGBColor.from_string(WHITE)
right_p = header_right.paragraphs[0]
right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
right_p.paragraph_format.space_after = Pt(0)
right_run = right_p.add_run("17 JULY 2026")
right_run.bold = True
right_run.font.size = Pt(9)
right_run.font.color.rgb = RGBColor.from_string(WHITE)

footer = section.footer
footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
footer_table.autofit = False
footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(footer_table, color=WHITE, size=0)
footer_left, footer_right = footer_table.rows[0].cells
set_cell_width(footer_left, 5.5)
set_cell_width(footer_right, 1.0)
footer_left_p = footer_left.paragraphs[0]
footer_left_p.paragraph_format.space_after = Pt(0)
footer_left_run = footer_left_p.add_run("Genetics-first screen • research interpretation only")
footer_left_run.font.size = Pt(8)
footer_left_run.font.color.rgb = RGBColor.from_string(MID_GRAY)
footer_right_p = footer_right.paragraphs[0]
footer_right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_right_p.paragraph_format.space_after = Pt(0)
footer_right_run = footer_right_p.add_run("Page ")
footer_right_run.font.size = Pt(8)
footer_right_run.font.color.rgb = RGBColor.from_string(MID_GRAY)
add_page_number(footer_right_p)

document.add_paragraph("Candidate genomic clusters", style="Title")
document.add_paragraph("Colleague email and reporting guide", style="Subtitle")
add_callout(
    document,
    "Bottom line: these are candidate genomic clusters. They are not nursing-home outbreaks or confirmed transmission clusters until Participant_id is linked to facility, ward/unit and timing information.",
    fill=PALE_AMBER,
    border="D6A400",
)

intro = document.add_paragraph()
intro.add_run("Analysis scope. ").bold = True
intro.add_run(
    f"The genetics-first screen used {summary['active_isolates']} current QC-passing isolates "
    f"from {summary['active_participants']} participants. It identified "
    f"{summary['high_priority_cluster_count']} high-priority candidate genomic clusters "
    f"involving {summary['high_priority_cluster_isolate_count']} isolates and "
    f"{summary['priority_counts']['High-priority candidate']} cross-participant high-priority pairs."
)

add_section_heading(document, "Copy-ready email", 1)
subject = document.add_paragraph()
subject_run = subject.add_run("Subject: Genetics-first shortlist of candidate genomic clusters")
subject_run.bold = True
subject_run.font.color.rgb = RGBColor.from_string(NAVY)

email_paragraphs = [
    "Hi [Name],",
    (
        f"I have screened the {summary['active_isolates']} QC-passing isolates for genetic similarity "
        "across different participants. I used reliable sequence type (ST) and the full 227-gene "
        "virulence-factor profile for initial screening, followed by pair-specific SNP distance as "
        "the more specific measure of strain-level relatedness."
    ),
    (
        f"Using the project's operational ≤25-SNP threshold, I identified "
        f"{summary['high_priority_cluster_count']} high-priority candidate genomic clusters involving "
        f"{summary['high_priority_cluster_isolate_count']} isolates. The workbook also retains "
        "moderate-priority pairs, possible related lineages, and pairs that share a lineage but were "
        "more than 25 SNPs apart on direct comparison."
    ),
    (
        "These should not yet be described as nursing-home outbreaks or confirmed transmission, "
        "because nursing-home and ward information was not available in the genomic dataset. "
        "The candidate tables are linked through Participant_id, Episode_ID and collection date, "
        "allowing you to add your nursing-home, ward/unit and timing variables and determine whether "
        "genetically related isolates also have a plausible epidemiological overlap."
    ),
    (
        "Resistance-gene results are available from a ResFinder analysis using calls at ≥80% identity "
        "and ≥80% coverage. I have summarised shared and differing resistance genes and the broad "
        "resistance classes represented. The near-ubiquitous mdf(A) gene was excluded from AMR-profile "
        "similarity calculations. These are genotypic predictions only; no phenotypic susceptibility "
        "or antibiogram data were available, so they should not be reported as confirmed clinical resistance."
    ),
    (
        "After you have linked the nursing-home information, we can review the strongest candidates "
        "together and agree how they should be described and reported."
    ),
    "Best wishes,\n[Your name]",
]
for text in email_paragraphs:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.add_run(text)

document.add_page_break()
add_section_heading(document, "High-priority clusters at a glance", 1)
document.add_paragraph(
    "All entries below meet the high-priority genetics rule across different participants: "
    "same reliable ST, an identical detected 227-gene VF profile, a preliminary core-SNP distance "
    "of ≤25, and a direct pairwise distance of ≤25 SNPs."
)

headers = ["Cluster", "Participant_id", "ST", "Isolates", "Direct SNPs", "Collection dates"]
widths = [0.70, 1.40, 0.55, 0.65, 0.85, 2.35]
table = document.add_table(rows=1, cols=len(headers))
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
set_table_borders(table)
header_row = table.rows[0]
set_repeat_table_header(header_row)
for index, (header_text, width) in enumerate(zip(headers, widths)):
    cell = header_row.cells[index]
    set_cell_width(cell, width)
    set_cell_fill(cell, NAVY)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(header_text)
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(WHITE)

for row_number, cluster in enumerate(clusters, start=1):
    values = [
        cluster["candidate_cluster_id"],
        cluster["Participant_ids"],
        cluster["ST"],
        cluster["n_isolates"],
        (
            cluster["direct_snp_min"]
            if cluster["direct_snp_min"] == cluster["direct_snp_max"]
            else f"{cluster['direct_snp_min']}–{cluster['direct_snp_max']}"
        ),
        f"{format_date(cluster['earliest_collection_date'])} to {format_date(cluster['latest_collection_date'])}",
    ]
    row = table.add_row()
    for index, (value, width) in enumerate(zip(values, widths)):
        cell = row.cells[index]
        set_cell_width(cell, width)
        set_cell_margins(cell, top=65, start=90, bottom=65, end=90)
        if row_number % 2 == 1:
            set_cell_fill(cell, PALE_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(str(value))
        run.font.size = Pt(8.3)
        if index == 0:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

document.add_paragraph(
    "Collection dates are shown descriptively; no arbitrary temporal cutoff was imposed. "
    "The colleague should assess whether participants were in the same nursing home or ward/unit "
    "during an epidemiologically plausible period."
)

add_section_heading(document, "What each measure contributes", 1)
add_labelled_paragraph(
    document,
    "ST — lineage screen. ",
    "The same reliable ST supports membership of the same broad bacterial lineage. It is useful for screening, but it is not strain-level proof and cannot establish transmission by itself.",
)
add_labelled_paragraph(
    document,
    "VF Jaccard — profile similarity. ",
    "This compares the presence/absence pattern across all 227 detected virulence-factor genes. A value of 1.00 means the profiles are identical; 0.95 means approximately 95% of the union of detected genes is shared. Similarity can support screening but is less specific than SNP distance.",
)
add_labelled_paragraph(
    document,
    "Preliminary core-SNP distance — shortlist. ",
    "This study-wide matrix was used to find promising pairs efficiently. It is a screen, not the final pair-specific result.",
)
add_labelled_paragraph(
    document,
    "Direct SNP distance — strain-level evidence. ",
    "The direct whole-assembly comparison is the primary pair-specific measure. The project uses ≤25 SNPs as its operational candidate threshold; that threshold supports close genomic relatedness but does not prove who transmitted to whom.",
)
add_labelled_paragraph(
    document,
    "Collection dates — temporal context. ",
    "Dates help the colleague assess whether a plausible overlap exists after facility and ward/unit data are linked. A long interval weakens a simple direct-transmission interpretation even when genomes are close.",
)
add_labelled_paragraph(
    document,
    "AMR genes — clinical description only. ",
    "ResFinder calls describe acquired resistance genes detected in the assemblies. AMR profiles do not determine cluster membership because genes may be gained or lost independently.",
)

add_section_heading(document, "How candidates were ranked", 1)
add_labelled_paragraph(
    document,
    "High-priority candidate. ",
    "Different participants; same reliable ST; identical 227-gene VF profile; preliminary core-SNP distance ≤25; and direct pairwise SNP distance ≤25.",
)
add_labelled_paragraph(
    document,
    "Moderate-priority candidate. ",
    "Different participants; same reliable ST; VF Jaccard ≥0.95 but not identical; preliminary SNP distance ≤25; direct validation pending or supporting information incomplete.",
)
add_labelled_paragraph(
    document,
    "Possible related lineage. ",
    "Different participants with the same ST and VF Jaccard 0.90–0.949, retained for checking against nursing-home, ward/unit and timing information.",
)
add_labelled_paragraph(
    document,
    "Same lineage, not the same strain. ",
    "The same ST and a similar or identical VF profile were present, but the direct distance was >25 SNPs. These should not be presented as transmission candidates under the project rule.",
)
add_callout(
    document,
    "Pairwise evidence remains visible in Candidate_pairs. This matters because cluster IDs are grouped from qualifying links, and pair-level distances should always be checked before making a cluster-wide statement.",
    fill=PALE_BLUE,
    border=BLUE,
)

add_section_heading(document, "How to report an individual cluster", 1)
add_labelled_paragraph(
    document,
    "Reusable template. ",
    "“Candidate cluster [ID] comprised [n] isolates from [n] participants. The isolates shared ST [number] and had [identical/highly similar] detected VF profiles. Direct pairwise distances were [range] SNPs, supporting genetically closely related isolates compatible with recent common ancestry. Facility, ward/unit and timing linkage is required before any transmission interpretation.”",
)
add_labelled_paragraph(
    document,
    "AMR sentence, if genes are present. ",
    "“The genomes contained [shared genes], associated with [broad resistance classes]. These are genotypic findings only; phenotypic susceptibility results were not available for confirmation.”",
)

cgc006 = next(cluster for cluster in clusters if cluster["candidate_cluster_id"] == "CGC-006")
add_labelled_paragraph(
    document,
    "Worked example — CGC-006. ",
    (
        f"“CGC-006 comprised {cgc006['n_isolates']} isolates from "
        f"{cgc006['n_participants']} participants, all ST{cgc006['ST']}, with identical detected "
        f"VF profiles and direct pairwise distances of {cgc006['direct_snp_min']}–"
        f"{cgc006['direct_snp_max']} SNPs. Excluding mdf(A), all members shared "
        f"{cgc006['shared_amr_genes_all_members_excluding_mdfA']}. The represented broad classes were "
        f"{cgc006['resistance_classes_present']}. This is a candidate genomic cluster requiring "
        "facility, ward/unit and temporal confirmation; the AMR findings are genotypic only.”"
    ),
)

add_section_heading(document, "AMR interpretation", 1)
add_labelled_paragraph(
    document,
    "Data available. ",
    "Existing ResFinder calls filtered at ≥80% identity and ≥80% coverage.",
)
add_labelled_paragraph(
    document,
    "How similarity was calculated. ",
    "AMR-profile Jaccard similarity excludes mdf(A) because it is near-ubiquitous and would otherwise inflate apparent similarity.",
)
add_labelled_paragraph(
    document,
    "What can be said. ",
    "“The genome contains resistance gene X, which is associated with resistance to antibiotic class Y; phenotypic susceptibility results were not available for confirmation.”",
)
add_labelled_paragraph(
    document,
    "What cannot be said. ",
    "Do not state that an isolate is clinically resistant to a particular antibiotic without MIC, disk-diffusion or other susceptibility results.",
)
add_labelled_paragraph(
    document,
    "Plasmids. ",
    "Shared resistance genes do not establish that isolates share the same plasmid. Plasmid-level sequence and structural confirmation would be required.",
)

add_section_heading(document, "What the colleague should do next", 1)
for number, text in enumerate(
    [
        "Join the nursing-home and ward/unit variables to both cluster and pair tables using Participant_id.",
        "Check whether participants in the same candidate cluster shared a facility or unit and whether their residence/care periods overlapped or followed a plausible sequence.",
        "Review Candidate_pairs, not only Cluster_summary, to confirm that the specific epidemiologically linked pair has the expected direct SNP distance.",
        "Use lower-priority pairs only as prompts for epidemiological review; do not upgrade them solely because they share a nursing home.",
        "Describe AMR genes separately from the evidence used to define genomic clusters.",
    ],
    start=1,
):
    add_labelled_paragraph(document, f"{number}. ", text)

add_section_heading(document, "Workbook map", 1)
add_labelled_paragraph(
    document,
    "Cluster_summary. ",
    "A concise view of the 10 high-priority candidate genomic clusters, with participants, episodes, isolate IDs, dates, ST, direct SNP ranges, one cross-participant pair count, and shared AMR genes/classes. Uniform high-priority criteria are stated once rather than repeated in every row.",
)
add_labelled_paragraph(
    document,
    "Candidate_pairs. ",
    "All 1,273 cross-participant screened pairs with preliminary SNP, VF, direct-validation and pair-level AMR evidence. This remains the key audit table for the details intentionally removed from Cluster_summary.",
)
add_labelled_paragraph(
    document,
    "Isolate_profiles. ",
    "One row for each of the 532 included isolates, with identifiers, QC, ST, VF count and ResFinder profile.",
)
add_labelled_paragraph(
    document,
    "Excluded_isolates. ",
    "The 24 emailed rows outside the current cohort: 22 failed current assembly QC with BadSize and two had no matching row in the current canonical assembly manifest.",
)
add_labelled_paragraph(
    document,
    "Interpretation_guide. ",
    "Plain-language definitions and recommended/avoided wording for ST, VF profiles, SNP distance, AMR, nursing-home linkage and plasmids.",
)

add_section_heading(document, "Preferred and avoided wording", 1)
wording_table = document.add_table(rows=1, cols=2)
wording_table.alignment = WD_TABLE_ALIGNMENT.CENTER
wording_table.autofit = False
set_table_borders(wording_table)
for index, (text, width) in enumerate((("Use", 3.25), ("Avoid", 3.25))):
    cell = wording_table.rows[0].cells[index]
    set_cell_width(cell, width)
    set_cell_fill(cell, DARK_BLUE)
    set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(WHITE)
set_repeat_table_header(wording_table.rows[0])

use_terms = [
    "candidate genomic cluster",
    "genetically closely related isolates",
    "compatible with recent common ancestry",
    "shared resistance-gene profile",
    "requires epidemiological confirmation",
]
avoid_terms = [
    "confirmed outbreak",
    "transmission occurred",
    "the same strain based only on ST or VF",
    "clinically resistant based only on gene detection",
    "shared plasmid without plasmid-level confirmation",
]
for index in range(max(len(use_terms), len(avoid_terms))):
    row = wording_table.add_row()
    for column, terms in enumerate((use_terms, avoid_terms)):
        cell = row.cells[column]
        set_cell_width(cell, 3.25)
        set_cell_margins(cell)
        set_cell_fill(cell, PALE_GREEN if column == 0 else PALE_RED)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(terms[index])
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(BLACK)

add_section_heading(document, "Limitations to state explicitly", 1)
add_labelled_paragraph(
    document,
    "No facility metadata in the genomic files. ",
    "Genetic clusters cannot be called nursing-home clusters until the colleague completes the Participant_id linkage.",
)
add_labelled_paragraph(
    document,
    "No phenotypic susceptibility data. ",
    "ResFinder results predict potential resistance mechanisms but do not confirm the clinical phenotype.",
)
add_labelled_paragraph(
    document,
    "No plasmid-level inference. ",
    "This analysis did not reconstruct or compare complete plasmid sequences.",
)
add_labelled_paragraph(
    document,
    "Operational SNP threshold. ",
    "The ≤25-SNP rule is the project threshold for candidate relatedness, not a universal biological boundary and not proof of direct transmission.",
)
add_labelled_paragraph(
    document,
    "Sampling limits. ",
    "Unsampled residents, staff, environmental sources or community intermediates may explain apparent links or gaps.",
)

add_callout(
    document,
    "Final reporting principle: keep three claims separate—genetic relatedness, epidemiological linkage, and predicted resistance.",
    fill=PALE_GREEN,
    border="70AD47",
)

document.save(OUTPUT_PATH)
print(f"Saved {OUTPUT_PATH}")
