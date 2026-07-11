#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(
    "/Users/Aamir/Desktop/rUTIs/outputs/manual-20260527-current-review/"
    "presentations/ruti-longitudinal-vf-pipeline-review-v3/output/"
    "Longitudinal_Urinary_Ecoli_VF_Pipeline_Detailed_Presenter_Script_2026-05-28.docx"
)
MD_OUT = OUT.with_suffix(".md")

BLUE = RGBColor(43, 108, 176)
GREEN = RGBColor(47, 133, 90)
ORANGE = RGBColor(217, 119, 6)
RUST = RGBColor(182, 90, 60)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(100, 116, 139)
LIGHT_BLUE = "E8F1FA"
LIGHT_GREEN = "EAF7EF"
LIGHT_ORANGE = "FFF4E6"
LIGHT_GREY = "F8FAFC"
LIGHT_RUST = "FBEDE8"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="D8DEE8", size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(int(w * 1440) for w in widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.15


def add_run(paragraph, text, bold=False, color=None, size=None, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    return run


def add_callout(doc, title, body, fill=LIGHT_BLUE, title_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [7.0])
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_fill(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, title, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, body)
    doc.add_paragraph()


def add_label_detail_table(doc, rows, widths=(1.55, 5.45), header_fill=None):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, list(widths))
    for i, (label, detail, fill, color) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_fill(c0, fill or LIGHT_GREY)
        if header_fill:
            set_cell_fill(c1, header_fill)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        add_run(p0, label, bold=True, color=color or BLUE)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        add_run(p1, detail)
    doc.add_paragraph()


def add_bullets(doc, bullets):
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, item)


slides = [
    {
        "n": 1,
        "time": "1:00",
        "title": "Virulence-factor profiling of longitudinal urinary E. coli isolates",
        "primary": "Reset the frame: this is a VF-pipeline and longitudinal-isolate review, not a recurrent-UTI-only talk.",
        "visual": "Point to the three question cards, then the four denominator tiles at the bottom.",
        "script": [
            "I want to frame this review around the current virulence-factor pipeline for longitudinal urinary E. coli isolates. The clinical labels matter, but the main object of study is the VF repertoire measured from repeated WGS-linked isolates.",
            "The three questions structure the talk. First: what VF repertoire is present in the current VF-ready dataset? Second: how stable are those profiles over time within the same resident? Third: once the pipeline is established, how can we overlay UTI versus Not_UTI status without overstating sparse clinical associations?",
            "The bottom numbers are the key anchors for the rest of the review: 556 VF-ready episodes, 227 VF gene columns, 394 longitudinal comparisons, and 17 UTI annotations in the VF/model-ready clinical overlay.",
        ],
        "emphasize": [
            "Say VF analysis first; clinical status overlay second.",
            "Make clear that the deck is about the current pipeline and what it can safely support.",
        ],
        "avoid": "Do not introduce this as a recurrent UTI findings deck or as a disease-association result.",
        "transition": "I’ll start with why repeated-isolate VF profiling gives us a different question from a single clinical-status comparison.",
    },
    {
        "n": 2,
        "time": "1:30",
        "title": "Longitudinal urinary E. coli VF analysis asks about repertoire and stability",
        "primary": "Longitudinal VF analysis separates persistent carriage, replacement, and gene-content change.",
        "visual": "Walk left to right: repeated isolates, three longitudinal patterns, clinical overlay as the later layer.",
        "script": [
            "For each isolate, the VF profile is a measured gene-content snapshot. If we only compare UTI and Not_UTI episodes cross-sectionally, we miss the temporal biology: a resident may carry a stable E. coli profile, may acquire a replacement lineage, or may show apparent VF gene gains or losses between visits.",
            "The middle column names the three patterns I want the listener to keep in mind: persistent lineage, replacement, and gene-content change. These are interpretive buckets, not mutually perfect mechanistic categories.",
            "The right-hand panel is deliberately labelled clinical overlay comes later. UTI_Status is useful, but in this analysis it is a cautious annotation layer placed on top of the genomic/VF pipeline rather than the whole thesis.",
        ],
        "emphasize": [
            "VF presence/absence is a genomic feature, not expression or activity.",
            "The longitudinal structure is the reason the dataset is interesting.",
        ],
        "avoid": "Do not say stable VF equals asymptomatic carriage or replacement equals new infection without supporting genomic context.",
        "transition": "The next slide shows what data asset feeds that longitudinal question.",
    },
    {
        "n": 3,
        "time": "1:30",
        "title": "VF-ready data link repeated urinary E. coli genomes to clinical episodes",
        "primary": "Define the unit of analysis and the VF-ready denominator.",
        "visual": "Use the metric tiles first, then the repeated isolate timeline.",
        "script": [
            "The analysis dataset is episode-level. Each row represents a selected urinary E. coli assembly linked to a clinical episode and translated into a binary VF profile.",
            "The VF-ready dataset contains 556 episode-level selected assemblies from 162 participants. Those assemblies are represented by 227 binary VF gene columns. The curation framework groups those genes into 32 modules, including 18 UPEC-candidate modules.",
            "The timeline underneath is schematic: one resident can contribute repeated Not_UTI and UTI-labelled episodes. The key point is that status labels can be overlaid on repeated isolates, but the molecular unit remains the episode-level isolate profile.",
        ],
        "emphasize": [
            "The 162 count is the VF-ready dataset participant count.",
            "Avoid displaying or discussing the unresolved total clinical participant count.",
        ],
        "avoid": "Do not imply that all 583 clinical episodes have VF-ready WGS evidence.",
        "transition": "Now I’ll go through the exact handoffs that take us from clinical classification to VF outputs.",
    },
    {
        "n": 4,
        "time": "2:30",
        "title": "Numbered clinical-to-VF pipeline with visible denominators",
        "primary": "Orient the colleague to each pipeline handoff and why counts change.",
        "visual": "Read the flow in numbered order; pause on the places where the denominator changes.",
        "script": [
            "This is the slide I would expect a new colleague to come back to. It shows the current pipeline with the count at each handoff, so they can understand what is being compared.",
            "We begin with 585 classified clinical episodes before primary manual exclusions. After primary clinical inclusion, 583 episodes remain: 18 UTI and 565 Not_UTI, with two excluded. At the WGS QC stage there are 1,291 assembly-level QC records, which is larger because assembler alternatives are still represented.",
            "Canonical selection brings this to 556 selected episode-level assemblies. Those become 556 VF presence/absence rows, and then 556 VF/model-ready episodes. In that model-ready clinical overlay there are 17 UTI and 539 Not_UTI episodes, meaning 27 clinical episodes do not have VF-ready evidence.",
            "The feature framework maps 227 VF gene columns into 32 modules, including 18 UPEC-candidate modules. The longitudinal layer then uses 394 consecutive within-resident comparisons from 144 participants. Finally, the clinical transition application identifies 11 Not_UTI to UTI clinical transitions, 10 of which are WGS/VF-linked, with one missing endpoint.",
        ],
        "emphasize": [
            "Different rows mean different units: clinical episodes, assembly QC records, selected assemblies, VF rows, longitudinal pairs.",
            "The main deck should not show disputed participant totals or ST totals.",
        ],
        "avoid": "Do not describe 1,291 assembly QC records as 1,291 unique isolates.",
        "transition": "Before we interpret VF figures, I want to define what the VF features are and what they are not.",
    },
    {
        "n": 5,
        "time": "2:00",
        "title": "VF features: genes, curated modules, and longitudinal similarity",
        "primary": "Separate raw gene detection, curated modules, and downstream outputs.",
        "visual": "Follow the three main boxes from gene matrix to modules to outputs, then explain the caution box.",
        "script": [
            "The raw VF representation is a binary presence/absence matrix: for each selected assembly, each of 227 VF gene columns is marked detected or not detected. That is a genomic detection framework, not a direct measure of transcription, protein production, or activity.",
            "The curated module layer helps make the feature space biologically interpretable. Rather than discussing hundreds of genes one by one, we can summarise systems such as adhesion, iron acquisition, secretion, toxin, capsule or other surface-related features, and unassigned genes.",
            "The outputs are burden, prevalence, Jaccard similarity, gain/loss, and exploratory models. The caution is important: 25.1 percent of the VF matrix is unassigned, so total burden should be interpreted separately from curated module or UPEC-candidate summaries.",
        ],
        "emphasize": [
            "Modules are curation units, not validated disease-causality scores.",
            "Jaccard similarity measures profile overlap, not strain identity by itself.",
        ],
        "avoid": "Do not let the audience hear 'module' as 'clinically validated pathway'.",
        "transition": "With the feature framework defined, the first evidence slide simply asks what VF genes are common.",
    },
    {
        "n": 6,
        "time": "1:30",
        "title": "Common VF genes describe the observed repertoire, not disease association",
        "primary": "The top-gene plot is descriptive repertoire evidence.",
        "visual": "Point to the ranked bars and the descriptive label.",
        "script": [
            "This figure ranks the most prevalent VF genes among the 556 VF/WGS-linked urinary E. coli isolates. It is a repertoire view: what genes are commonly detected in this dataset?",
            "You can use it to orient the audience to the feature space: common fimbrial or adhesion-associated genes, iron acquisition systems, and other VFDB-derived features appear across many isolates. But this is not an association result. A gene being common does not mean it distinguishes UTI from Not_UTI or causes symptoms.",
            "If someone asks why this matters, the answer is that it defines the background VF landscape before we start discussing longitudinal stability or clinical overlays.",
        ],
        "emphasize": [
            "Descriptive prevalence comes before association testing.",
            "The y-axis ranking is about frequency in the VF-ready dataset.",
        ],
        "avoid": "Do not call any prevalent gene a UTI marker from this slide.",
        "transition": "The next slide groups those individual genes into biological modules to make interpretation easier.",
    },
    {
        "n": 7,
        "time": "1:30",
        "title": "Curated modules organise genes into interpretable biological systems",
        "primary": "Modules reduce complexity while preserving biological interpretability.",
        "visual": "Use the bar heights to show which module categories contain more genes.",
        "script": [
            "This slide shows the module framework. The goal is not to replace gene-level information, but to give the analysis a more interpretable vocabulary.",
            "A new colleague can use this figure to understand which biological systems are represented in the VF matrix. Some modules contain many genes, so they will naturally contribute more to burden summaries, while smaller modules may be biologically meaningful but numerically less dominant.",
            "The important phrasing is curated framework. These modules help us organise the analysis; they are not validated predictors of symptomatic infection.",
        ],
        "emphasize": [
            "Modules are helpful for biological discussion and handover.",
            "Gene count per module affects how burden-like summaries behave.",
        ],
        "avoid": "Do not compare module size as though it directly measures importance.",
        "transition": "Once the repertoire and modules are defined, the central longitudinal question is how much they change over time.",
    },
    {
        "n": 8,
        "time": "2:00",
        "title": "Most repeated within-resident VF profiles are highly stable",
        "primary": "The central longitudinal result: VF profiles are often conserved within residents.",
        "visual": "Point to the mass of comparisons near Jaccard 1.000 and the central finding box.",
        "script": [
            "This is one of the most important slides in the deck. Across 394 consecutive within-resident comparisons from 144 participants, the median within-resident VF Jaccard similarity is 1.000. In other words, for the typical consecutive pair, the detected VF gene set is identical.",
            "The additional headline is that 62.4 percent of consecutive comparisons show no VF change. That tells us that the dominant pattern in this dataset is profile stability, not repeated major VF gain or loss.",
            "The careful interpretation is that VF stability supports persistent-profile explanations, but it does not prove same-strain persistence on its own. For that, we need SNP distance, sequence type, and broader genomic context.",
        ],
        "emphasize": [
            "Median Jaccard similarity: 1.000.",
            "62.4 percent with no VF change.",
            "This is descriptive longitudinal evidence.",
        ],
        "avoid": "Do not claim that stable VF profile proves the same strain persisted.",
        "transition": "The next slide asks where change does appear and how we should treat it.",
    },
    {
        "n": 9,
        "time": "2:00",
        "title": "When profiles change, gain/loss summaries flag candidates for follow-up",
        "primary": "VF change events are leads, not automatic biological conclusions.",
        "visual": "Explain the gain/loss distributions and point to outlier behaviour.",
        "script": [
            "This figure focuses on consecutive pairs where the VF profile does change. It summarises apparent gene gains and losses between repeated isolates.",
            "Most comparisons are close to zero, consistent with the stability shown on the previous slide. But the tails matter. Pairs with many apparent gains or losses are candidates for follow-up, because they may represent strain replacement, changes in assembly or gene-calling confidence, or true gene-content change.",
            "The correct language here is flagging and prioritisation. This slide helps us identify which cases deserve closer review; it does not by itself adjudicate mechanism.",
        ],
        "emphasize": [
            "Gain/loss summaries should trigger context review.",
            "Outliers are useful even if they are not immediately mechanistic.",
        ],
        "avoid": "Do not call every gain/loss event within-host evolution.",
        "transition": "To interpret those changes, we need lineage context.",
    },
    {
        "n": 10,
        "time": "1:30",
        "title": "Sequence-type consistency helps interpret VF stability",
        "primary": "Lineage context helps distinguish persistence-like from replacement-like comparisons.",
        "visual": "Contrast same-ST and different-ST comparisons without overreading the plot.",
        "script": [
            "This diagnostic plot asks whether VF similarity behaves differently when consecutive isolates share the same sequence type versus when they do not. It provides lineage context for the stability result.",
            "The broad expectation is that same-ST pairs should often have more similar VF profiles, while different-ST pairs are more likely to reflect replacement-like events or broader gene-content differences. But ST is still a coarse label. It helps triage interpretation; it does not prove persistence alone.",
            "This is why the pipeline combines profile similarity with SNP distance and case-level interpretation when discussing transitions.",
        ],
        "emphasize": [
            "ST is a context variable, not a final answer.",
            "This slide prepares the audience for cautious clinical interpretation.",
        ],
        "avoid": "Do not equate same ST with same strain.",
        "transition": "Only after the VF and lineage context are established do I introduce clinical-status comparison.",
    },
    {
        "n": 11,
        "time": "2:00",
        "title": "UTI status overlay remains exploratory under the sparse denominator",
        "primary": "No global VF association remains significant after FDR correction.",
        "visual": "Explain nominal screening versus participant-aware model evidence; point to the safe-claim box.",
        "script": [
            "This is the first clinical-status application. The model-ready clinical overlay contains 556 episodes: 17 UTI and 539 Not_UTI. That imbalance is the central reason for caution.",
            "The figure compares exploratory gene screening with participant-aware modelling. Some genes may look interesting nominally, but the current safe conclusion is that no global VF association remains significant after FDR correction.",
            "So this slide should be used for hypothesis generation and for explaining uncertainty. It is not a claim that the pipeline has found a robust VF marker of symptomatic UTI.",
        ],
        "emphasize": [
            "17 UTI versus 539 Not_UTI in the VF/model-ready dataset.",
            "No global VF association survives FDR correction.",
            "Clinical overlay is exploratory.",
        ],
        "avoid": "Do not mention older ASB-versus-UTI conclusions or significant lpf claims as current results.",
        "transition": "The next slide makes the clinical overlay concrete with one participant transition.",
    },
    {
        "n": 12,
        "time": "2:00",
        "title": "Participant 20026 illustrates stable VF profile despite symptom emergence",
        "primary": "A concrete transition shows symptoms can emerge without detectable VF gain.",
        "visual": "Walk across the timeline, then the four green evidence tiles, then the two summary boxes.",
        "script": [
            "Participant 20026 is a useful worked example because it links the abstract pipeline to a real longitudinal transition. The participant moves from Not_UTI to UTI over 42 days.",
            "The genomic and VF evidence shows 5 SNPs and a stable VF profile: zero VF genes gained and zero lost, with stable VF modules and strong same-strain evidence. That pattern supports the idea that symptom emergence does not always require a detectable VF repertoire change.",
            "At the casebook level, there are 11 clinical Not_UTI to UTI transitions, 10 WGS/VF-linked transitions, and one missing endpoint. Four transitions show same-strain stable profiles, while three are consistent with strain replacement. These buckets organise evidence. They do not prove mechanism.",
            "The careful interpretation is that low SNP distance plus stable VF profile supports host-state, expression, regulation, inoculum, or other unmeasured explanations, but it does not prove any one of them.",
        ],
        "emphasize": [
            "Not_UTI to UTI over 42 days.",
            "5 SNPs, stable VF profile, zero gains/losses.",
            "Support, not proof, of host-state or unmeasured-regulation hypotheses.",
        ],
        "avoid": "Do not say stable VF proves host mechanism or symptom causality.",
        "transition": "I’ll close the spoken section by separating what is established from what remains uncertain.",
    },
    {
        "n": 13,
        "time": "1:00",
        "title": "What this VF-first review establishes and what remains uncertain",
        "primary": "Close with evidence, boundary, and next steps.",
        "visual": "Use the three columns as the closing structure.",
        "script": [
            "The evidence column is the firm ground: the current pipeline produces a coherent VF-ready longitudinal dataset with 556 isolate profiles, 227 VF gene columns, 32 modules, and 394 consecutive within-resident comparisons.",
            "The boundary column is equally important. VF presence/absence does not measure expression or activity, and clinical-status association remains exploratory because the UTI denominator is sparse.",
            "The next-step column is where a new colleague can contribute: lineage-aware longitudinal follow-up, expression or regulation hypotheses, and targeted review of transition case studies. The main discussion question is what additional evidence would distinguish stable carriage with host-state change from unmeasured bacterial regulation or replacement.",
        ],
        "emphasize": [
            "End on usable pipeline plus honest uncertainty.",
            "Invite discussion around next evidence, not around overclaiming current associations.",
        ],
        "avoid": "Do not close as though the VF analysis has identified definitive UTI drivers.",
        "transition": "From here, the appendix is available for operational handover and detailed questions.",
    },
]

appendix = [
    {
        "n": 14,
        "title": "Practical map: where to enter and rerun the VF pipeline",
        "use": "Open when the colleague asks how to navigate the project or reproduce outputs.",
        "script": [
            "This is the handover slide. I would use it after the main talk, not during the core narrative unless the audience asks about reproducibility.",
            "The key message is that the colleague should start from the canonical VF dataset and diagnostics, then use the pipeline documentation and final figure pack to trace outputs. The final figure pack script is the entry point for current final figures; the VF and longitudinal outputs are the sources of truth for pipeline-specific questions.",
        ],
        "safe": "These are current navigation anchors and rerun order, not new analysis results.",
        "avoid": "Do not ask the colleague to start from old ASB scripts or legacy figures.",
    },
    {
        "n": 15,
        "title": "Clinical phenotype denominator and definition",
        "use": "Open when someone asks how UTI and Not_UTI were defined.",
        "script": [
            "This appendix slide explains the clinical annotation layer. It shows the validated clinical denominator: 583 primary clinical episodes, with 18 UTI and 565 Not_UTI.",
            "The clinical rule is catheter-aware and symptom-supported, with culture support. The important framing is that Not_UTI is heterogeneous: it can include culture-supported bacteriuria without compatible symptoms and other non-UTI states. In this VF-first deck, this status layer is used later as an exploratory annotation.",
        ],
        "safe": "UTI_Status is current UTI versus Not_UTI; it is not the older ASB-versus-UTI storyline.",
        "avoid": "Do not treat Not_UTI as a single biologically uniform state.",
    },
    {
        "n": 16,
        "title": "Module prevalence by status",
        "use": "Open for questions about whether module prevalence differs by UTI_Status.",
        "script": [
            "This plot gives module prevalence by clinical annotation. It is useful for visual pattern recognition, but the label is exploratory for a reason.",
            "The safe interpretation is that modules can be compared descriptively across UTI and Not_UTI episodes, but sparse UTI counts, repeated measures, and lineage structure prevent strong clinical conclusions from this slide alone.",
        ],
        "safe": "Exploratory clinical annotation only.",
        "avoid": "Do not call module differences validated UTI signatures.",
    },
    {
        "n": 17,
        "title": "Full transition mechanism casebook",
        "use": "Open when someone wants to inspect all Not_UTI to UTI transition categories.",
        "script": [
            "This is the aggregate transition casebook. It summarises clinical Not_UTI to UTI transitions and the genomic/VF evidence attached to them.",
            "Use it to show that transitions are heterogeneous. Some are consistent with same-strain stable-profile transitions, some with replacement, and some remain less certain. The purpose is evidence organisation, not mechanism proof.",
        ],
        "safe": "11 clinical transitions, 10 WGS/VF-linked, with stable-profile and replacement-consistent categories.",
        "avoid": "Do not imply that every transition has a single settled biological mechanism.",
    },
    {
        "n": 18,
        "title": "Stable strain and changing clinical state",
        "use": "Open when someone asks how low SNP distance and stable VF relate to symptom emergence.",
        "script": [
            "This slide expands the idea shown in participant 20026. It illustrates cases where low SNP distance and stable measured VF profile coexist with changing clinical state.",
            "The interpretive value is that symptoms can change even when the measured VF repertoire does not. That points to host context, regulation, expression, inoculum, tissue state, or other unmeasured factors as plausible follow-up hypotheses.",
        ],
        "safe": "Stable measured profiles support alternative hypotheses; they do not prove them.",
        "avoid": "Do not say host state is proven.",
    },
    {
        "n": 19,
        "title": "Population-level robustness boundary",
        "use": "Open for statistical robustness questions about the global VF overlay.",
        "script": [
            "This slide reinforces the population-level boundary. The analysis is underpowered for robust clinical-status association because there are only 17 VF-ready UTI episodes.",
            "Use the slide to explain why the deck presents population-level VF findings as exploratory and why no global VF signal should be treated as confirmatory after FDR correction.",
        ],
        "safe": "No adjusted VF association is confirmatory with only 17 VF-ready UTI episodes.",
        "avoid": "Do not overinterpret nominal effects.",
    },
    {
        "n": 20,
        "title": "Lineage structure is an interpretation check",
        "use": "Open when someone asks about population structure or confounding by lineage.",
        "script": [
            "This PCoA is a diagnostic view of VF profile structure by sequence type. It reminds us that VF profiles are not randomly distributed across isolates; they are shaped by lineage.",
            "That matters because a clinical-status comparison can be confounded by lineage composition. The correct use is diagnostic: lineage structure needs to be considered before clinical interpretation.",
        ],
        "safe": "Diagnostic/descriptive lineage context.",
        "avoid": "Do not use the PCoA as a disease-association result.",
    },
    {
        "n": 21,
        "title": "AMR backup only: exploratory transition-level ResFinder context",
        "use": "Open only if asked about AMR or accessory/plasmid changes.",
        "script": [
            "This is intentionally an appendix-only backup. ResFinder and accessory/plasmid context can help describe individual transitions, but the current deck does not present a completed global VF-plus-AMR association analysis.",
            "If asked, say that AMR is useful context for transition-level interpretation, but it should not be elevated to a primary result in this VF pipeline review.",
        ],
        "safe": "Exploratory transition-level context only.",
        "avoid": "Do not claim global AMR association or integrate AMR as a primary spoken result.",
    },
    {
        "n": 22,
        "title": "Sources of truth and lookup resources for questions",
        "use": "Open at the end or during Q&A when someone asks where a number, figure, or caveat comes from.",
        "script": [
            "This is the evidence registry. It points the colleague to the count validation files, denominator flow, VF diagnostics, WGS registry, longitudinal summaries, module notes, casebook, and phenotype explanation.",
            "The most important handover message is that questions should be answered from current sources of truth rather than from older exploratory slides or superseded analyses.",
        ],
        "safe": "Use these files to verify provenance and avoid stale claims.",
        "avoid": "Do not use backup diagnostics to create new claims during the presentation.",
    },
]

avoid_saying = [
    "This is a recurrent UTI presentation.",
    "ASB versus UTI is the current primary comparison.",
    "lpf is significantly associated with UTI in the current final analysis.",
    "Stable VF profile proves host-state mechanism.",
    "No VF gene change means nothing bacterial changed.",
    "AMR is a main global association result in this deck.",
    "Same sequence type proves same strain.",
]


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(
        p,
        "Detailed Presenter Script",
        bold=True,
        color=BLUE,
        size=24,
    )
    p = doc.add_paragraph()
    add_run(
        p,
        "Virulence-factor profiling of longitudinal urinary E. coli isolates",
        bold=True,
        color=INK,
        size=16,
    )
    p = doc.add_paragraph()
    add_run(
        p,
        "Current VF pipeline, descriptive longitudinal findings, and cautious clinical-status overlay",
        color=MUTED,
        size=12,
    )
    doc.add_paragraph()
    add_label_detail_table(
        doc,
        [
            ("Purpose", "Support a 22-minute scientific review plus discussion for a colleague who needs biological interpretation and project-navigation context.", LIGHT_BLUE, BLUE),
            ("Audience", "New team colleague with mixed clinical and bioinformatics/R orientation needs.", LIGHT_GREY, BLUE),
            ("Framing", "VF analysis of longitudinal urinary E. coli isolates first; UTI versus Not_UTI only as an exploratory clinical annotation.", LIGHT_GREEN, GREEN),
            ("Core timing", "Slides 1-13 are the spoken narrative. Slides 14-22 are appendix/Q&A backup.", LIGHT_ORANGE, ORANGE),
        ],
    )
    add_callout(
        doc,
        "One-sentence thesis",
        "The current pipeline produces a coherent VF-ready longitudinal isolate dataset; repeated within-resident VF profiles are commonly stable; clinical-status associations remain exploratory under a sparse UTI denominator.",
        fill=LIGHT_GREEN,
        title_color=GREEN,
    )
    doc.add_page_break()


def add_slide_section(doc, item, appendix_mode=False):
    doc.add_heading(f"Slide {item['n']}. {item['title']}", level=1)
    if appendix_mode:
        add_label_detail_table(
            doc,
            [
                ("When to open", item["use"], LIGHT_BLUE, BLUE),
                ("Safe claim", item["safe"], LIGHT_GREEN, GREEN),
                ("Avoid", item["avoid"], LIGHT_ORANGE, ORANGE),
            ],
        )
        doc.add_heading("Speaker Script", level=2)
        for para in item["script"]:
            doc.add_paragraph(para)
        return

    add_label_detail_table(
        doc,
        [
            ("Target time", item["time"], LIGHT_GREY, BLUE),
            ("Primary message", item["primary"], LIGHT_BLUE, BLUE),
            ("Visual walkthrough", item["visual"], LIGHT_GREEN, GREEN),
        ],
    )
    doc.add_heading("Speaker Script", level=2)
    for para in item["script"]:
        doc.add_paragraph(para)
    doc.add_heading("Emphasise", level=3)
    add_bullets(doc, item["emphasize"])
    add_label_detail_table(
        doc,
        [
            ("Avoid", item["avoid"], LIGHT_ORANGE, ORANGE),
            ("Transition", item["transition"], LIGHT_GREY, BLUE),
        ],
    )


def write_markdown():
    lines = [
        "# Detailed Presenter Script",
        "",
        "**Deck:** Virulence-factor profiling of longitudinal urinary *E. coli* isolates",
        "",
        "**Framing:** VF analysis of longitudinal urinary *E. coli* isolates first; UTI versus `Not_UTI` only as an exploratory clinical annotation.",
        "",
        "## Avoid Saying",
    ]
    for item in avoid_saying:
        lines.append(f"- {item}")
    lines.append("")
    lines.append("## Spoken Slides")
    for slide in slides:
        lines.extend(
            [
                "",
                f"### Slide {slide['n']}. {slide['title']} ({slide['time']})",
                f"**Primary message:** {slide['primary']}",
                f"**Visual walkthrough:** {slide['visual']}",
                "",
                "**Speaker script:**",
            ]
        )
        for para in slide["script"]:
            lines.append(f"\n{para}")
        lines.append("")
        lines.append("**Emphasise:**")
        for bullet in slide["emphasize"]:
            lines.append(f"- {bullet}")
        lines.append(f"\n**Avoid:** {slide['avoid']}")
        lines.append(f"\n**Transition:** {slide['transition']}")
    lines.append("")
    lines.append("## Appendix / Q&A Slides")
    for slide in appendix:
        lines.extend(
            [
                "",
                f"### Slide {slide['n']}. {slide['title']}",
                f"**When to open:** {slide['use']}",
                f"**Safe claim:** {slide['safe']}",
                f"**Avoid:** {slide['avoid']}",
                "",
                "**Speaker script:**",
            ]
        )
        for para in slide["script"]:
            lines.append(f"\n{para}")
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


def build():
    doc = Document()
    style_doc(doc)
    add_title_page(doc)

    doc.add_heading("How To Use This Script", level=1)
    doc.add_paragraph(
        "Read the speaker script as a rehearsal version, not a word-for-word obligation. "
        "The primary message should be said nearly exactly; the paragraphs can be shortened live."
    )
    add_callout(
        doc,
        "Tone",
        "Use VF-first language. Bring in UTI versus Not_UTI only when the deck reaches the clinical annotation section.",
        fill=LIGHT_BLUE,
        title_color=BLUE,
    )
    doc.add_heading("Avoid Saying", level=2)
    add_bullets(doc, avoid_saying)
    doc.add_page_break()

    doc.add_heading("Slides 1-13: Spoken Narrative", level=1)
    for slide in slides:
        add_slide_section(doc, slide)

    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    doc.add_heading("Slides 14-22: Appendix / Q&A Prompts", level=1)
    for slide in appendix:
        add_slide_section(doc, slide, appendix_mode=True)

    doc.add_heading("Final 30-Second Close", level=1)
    doc.add_paragraph(
        "If I had to compress the whole deck into one message, I would say: this project now has a "
        "current, navigable VF pipeline for longitudinal urinary E. coli isolates. The strongest "
        "descriptive signal is within-resident VF profile stability. The clinical overlay is useful "
        "for generating hypotheses, especially around symptom emergence without detectable VF gain, "
        "but it remains exploratory because the UTI denominator is sparse."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    write_markdown()
    print(OUT)
    print(MD_OUT)


if __name__ == "__main__":
    build()
