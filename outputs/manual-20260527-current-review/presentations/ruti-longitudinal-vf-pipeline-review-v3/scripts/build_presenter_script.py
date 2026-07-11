#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path("/Users/Aamir/Desktop/rUTIs/outputs/manual-20260527-current-review/presentations/ruti-longitudinal-vf-pipeline-review-v3/output/Longitudinal_Urinary_Ecoli_VF_Pipeline_Presenter_Script_2026-05-27.docx")

BLUE = RGBColor(43, 108, 176)
GREEN = RGBColor(47, 133, 90)
ORANGE = RGBColor(217, 119, 6)
RUST = RGBColor(182, 90, 60)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(100, 116, 139)
LIGHT_BLUE = "E8F1FA"
LIGHT_GREEN = "EAF7EF"
LIGHT_ORANGE = "FFF4E6"
LIGHT_GREY = "F8FAFC"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_run(paragraph, text, bold=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    return run


def add_label_para(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, f"{label}: ", bold=True, color=BLUE)
    add_run(p, text)


def add_callout(doc, title, body, fill=LIGHT_BLUE, title_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_fill(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    add_run(p, title, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, body)
    doc.add_paragraph()


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text)


slides = [
    {
        "n": 1,
        "title": "Virulence-factor profiling of longitudinal urinary E. coli isolates",
        "time": "1 min",
        "message": "Open by resetting the centre of gravity: this is a VF pipeline and longitudinal isolate review, with clinical status as a later annotation.",
        "script": "I want to frame this as a virulence-factor analysis of longitudinal urinary E. coli isolates. The clinical labels are important, but they are not the whole story. The first question is what VF repertoire we observe in the current WGS-linked dataset. The second is how stable those profiles are when the same resident contributes repeated isolates. The third is how far we can cautiously overlay UTI versus Not_UTI status, given the sparse UTI denominator.",
        "transition": "I will start with why the longitudinal VF framing is more informative than a single clinical-status comparison.",
        "caution": "Do not open as a recurrent-UTI-only talk; make the VF pipeline the subject.",
    },
    {
        "n": 2,
        "title": "Longitudinal VF analysis asks about repertoire and stability",
        "time": "1.5 min",
        "message": "Explain the biological reason to track VF profiles over time.",
        "script": "For a urinary E. coli isolate, a VF profile is a measured gene-content snapshot. If we only compare status groups cross-sectionally, we miss the longitudinal question: does a resident carry a stable VF profile over time, do they acquire a replacement strain, or do specific VF genes appear or disappear between isolates? This slide separates those possibilities. UTI status remains useful, but it should come in after we understand the measured VF landscape and its stability.",
        "transition": "That leads to the actual data asset: the VF-ready repeated-isolate dataset.",
        "caution": "Do not imply that VF presence means expression or activity.",
    },
    {
        "n": 3,
        "title": "A VF-ready longitudinal isolate dataset",
        "time": "1.5 min",
        "message": "State the VF-ready denominator and unit of analysis.",
        "script": "The analysis dataset contains 556 episode-level selected assemblies from 162 participants. Each row is a urinary E. coli isolate profile linked to a clinical episode and to a binary VF gene matrix. The main molecular feature space has 227 VF gene columns, then those are grouped into 32 curated modules, including 18 UPEC-candidate modules. I am deliberately showing the VF-ready participant count here, not the disputed full clinical participant total.",
        "transition": "Now I will walk through the handoffs that produce this table.",
        "caution": "Use the 162 participant count only for the VF-ready dataset.",
    },
    {
        "n": 4,
        "title": "Numbered clinical-to-VF pipeline with denominators",
        "time": "2.5 min",
        "message": "Show the current pipeline and the numbers retained at each handoff.",
        "script": "This is the main orientation slide for the pipeline. We start with 585 classified clinical episodes before primary manual exclusions. The primary clinical inclusion set is 583 episodes: 18 UTI and 565 Not_UTI, with two exclusions. The assembly QC level has 1,291 records because assembler alternatives are still explicit rows. Canonical selection reduces this to 556 selected episode-level assemblies. Those become 556 VF presence/absence rows and then 556 VF/model-ready episodes: 17 UTI and 539 Not_UTI. In parallel, the feature framework maps 227 VF gene columns into 32 modules, and the longitudinal analysis uses 394 consecutive within-resident comparisons from 144 participants.",
        "transition": "Before interpreting figures, I want to make clear what the VF features actually mean.",
        "caution": "Emphasise that the 1,291 QC records are assembly-level rows, not independent episodes.",
    },
    {
        "n": 5,
        "title": "How VF features are represented",
        "time": "2 min",
        "message": "Separate raw gene matrix, modules, and downstream outputs.",
        "script": "The raw input for VF analysis is a binary gene matrix: a gene is detected or not detected in each isolate. The module layer is a curation layer, not a causal score. It helps us discuss adhesion, iron acquisition, secretion, toxins, capsule or surface structures, and unassigned genes in a consistent way. One important caveat is that unassigned genes are 25.1 percent of the VF matrix, so total VF burden should not be treated as equivalent to curated or UPEC-candidate burden.",
        "transition": "With that feature framework in mind, the first descriptive result is the common VF repertoire.",
        "caution": "Do not call modules validated UTI predictors.",
    },
    {
        "n": 6,
        "title": "Common VF genes describe the observed repertoire",
        "time": "1.5 min",
        "message": "Present top-gene prevalence as descriptive repertoire, not association.",
        "script": "This figure is descriptive. It ranks the most prevalent VF genes among the 556 VF/WGS-linked urinary E. coli isolates. The main point is that the cohort contains a broad and common VF repertoire, including adhesion and fimbrial genes, iron-acquisition systems, and unassigned VFDB-derived genes. It is not testing whether any of these genes cause symptoms or distinguish UTI from Not_UTI.",
        "transition": "The next slide collapses the gene list into curated biological modules.",
        "caution": "Prevalence ranking is not an association test.",
    },
    {
        "n": 7,
        "title": "Curated modules organise genes into biological systems",
        "time": "1.5 min",
        "message": "Explain module curation as interpretation infrastructure.",
        "script": "Here the same gene space is organised by curated biological module. This is useful because individual genes can be hard to reason through in a live review. The framework lets a colleague quickly see which systems dominate the matrix, such as adhesins, iron acquisition, secretion systems, and prophage or mobile elements. Again, the module layer is navigation and curation, not a validated disease score.",
        "transition": "Once we know the feature space, the central longitudinal question is whether those profiles change over time.",
        "caution": "Keep the word 'curated' distinct from 'validated'.",
    },
    {
        "n": 8,
        "title": "Most repeated within-resident VF profiles are highly stable",
        "time": "2 min",
        "message": "Present the core longitudinal result.",
        "script": "This is the central longitudinal result. Across 394 consecutive within-resident isolate-pair comparisons from 144 participants, the median Jaccard similarity is 1.000, and 62.4 percent of pairs have no VF change. So the dominant pattern is VF profile conservation across repeated urinary E. coli isolates from the same resident. That does not prove same-strain persistence by itself, but it tells us that large VF gene-content shifts are not the usual pattern in this dataset.",
        "transition": "The next view shows where gain and loss events do appear.",
        "caution": "High VF similarity supports stability but does not prove same strain without SNP or lineage context.",
    },
    {
        "n": 9,
        "title": "Gain/loss summaries flag follow-up candidates",
        "time": "2 min",
        "message": "Use gain/loss plots as descriptive leads.",
        "script": "This plot separates VF genes gained and lost between repeated isolates. Most consecutive comparisons sit near zero, but there are outliers where many genes appear to change. Those outliers are useful because they point to candidates for closer review, but they need context: gain or loss can represent strain replacement, assembly or calling differences, or true gene-content change.",
        "transition": "That is why the next slide adds lineage context.",
        "caution": "Do not call every gain/loss event biological evolution.",
    },
    {
        "n": 10,
        "title": "Sequence-type consistency helps interpret VF stability",
        "time": "1.5 min",
        "message": "Place lineage context before status interpretation.",
        "script": "This diagnostic view asks whether repeated isolates with the same sequence type also tend to have high VF similarity. The point is interpretive: lineage context helps separate persistent-lineage stability from replacement-like changes. ST agreement is not a full same-strain proof, but it is an important guardrail before we start discussing clinical-status associations.",
        "transition": "Only now do I bring in the clinical-status overlay.",
        "caution": "Do not use ST agreement alone as proof of same strain.",
    },
    {
        "n": 11,
        "title": "UTI status overlay remains exploratory",
        "time": "2 min",
        "message": "State the sparse denominator and FDR result clearly.",
        "script": "This is the first clinical annotation application. In the VF/model-ready dataset, there are 556 episodes: 17 UTI and 539 Not_UTI. The visual bridges nominal screening evidence to participant-aware modelling. The safe conclusion is that no global VF association remains significant after FDR correction. So clinical-status patterns can generate hypotheses, but they are not confirmatory results.",
        "transition": "The final applied example makes that caution concrete in one participant transition.",
        "caution": "Do not resurrect older significant-gene claims or imply a robust global VF association.",
    },
    {
        "n": 12,
        "title": "Participant 20026: stable VF profile with symptom emergence",
        "time": "2 min",
        "message": "Use one transition to make the applied clinical question concrete.",
        "script": "Participant 20026 transitions from Not_UTI to UTI over 42 days. The paired genomes show 5 SNPs and a stable VF/module profile: zero VF genes gained and zero lost. This supports the idea that symptom emergence does not always require a detectable VF gain. Across the casebook, there are 11 clinical Not_UTI to UTI transitions, 10 WGS/VF-linked transitions, and 1 missing endpoint. Of these, 4 are same-strain stable-profile transitions, while 3 are consistent with strain replacement. The key wording is support, not proof: stable VF plus low SNP distance leaves host state, expression, regulation, inoculum, and other unmeasured factors on the table.",
        "transition": "I will close by summarising what the VF-first review establishes.",
        "caution": "Do not say stable VF establishes host mechanism.",
    },
    {
        "n": 13,
        "title": "Takeaways and handover transition",
        "time": "1 min",
        "message": "Close with evidence, boundary, and next steps.",
        "script": "The take-home is threefold. First, the current pipeline produces a usable VF-ready longitudinal dataset. Second, the dominant within-resident pattern is VF profile stability, with gain/loss events acting as follow-up leads. Third, clinical-status overlays are useful but exploratory, especially with only 17 VF-ready UTI episodes. The next practical step for a colleague is to understand the canonical outputs and where each figure comes from.",
        "transition": "If questions become operational, I would open the appendix handover slide.",
        "caution": "Keep the close focused on VF pipeline value and uncertainty, not a disease-causality claim.",
    },
]

appendix = [
    (14, "Practical handover and rerun order", "Open this when the colleague asks where to start. Safe claim: these are the canonical current outputs and the final-figure entry point."),
    (15, "Clinical phenotype denominator and definition", "Open this for questions about how UTI_Status is defined. Safe claim: clinical status is the annotation layer, not the main VF result."),
    (16, "Module prevalence by status", "Open this for exploratory status-context questions. Safe claim: descriptive only; repeated measures and lineage structure limit interpretation."),
    (17, "Full transition mechanism casebook", "Open this for Not_UTI -> UTI transition questions. Safe claim: 11 clinical transitions, 10 WGS/VF-linked, one missing endpoint."),
    (18, "Stable strain and clinical-state interpretation", "Open this for same-strain/stable-profile questions. Safe claim: supports host-state or regulation hypotheses, but does not prove mechanism."),
    (19, "Population-level robustness boundary", "Open this if someone asks whether global VF signals are significant. Safe claim: no adjusted association is confirmatory."),
    (20, "Lineage structure diagnostic", "Open this for ST/lineage questions. Safe claim: lineage structure must be considered before interpreting status associations."),
    (21, "AMR backup only", "Open only if asked about ResFinder context. Safe claim: exploratory transition-level context; not a global VF-plus-AMR association analysis."),
    (22, "Evidence registry and lookup resources", "Open for provenance. Safe claim: these files are where the current counts, captions, diagnostics, and pipeline documentation live."),
]


def build():
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Presenter Script")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(26)
    run.font.color.rgb = INK
    run.bold = True

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(16)
    add_run(sub, "Virulence-factor profiling of longitudinal urinary E. coli isolates: current pipeline and clinical context", color=MUTED, size=13)

    meta = doc.add_table(rows=4, cols=2)
    meta.autofit = False
    set_table_width(meta, [1.55, 4.95])
    for row in meta.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=80, bottom=80, start=120, end=120)
    for i, (k, v) in enumerate([
        ("Purpose", "Rehearsal script for a 22-minute VF-first scientific review plus discussion."),
        ("Audience", "New colleague who needs biological interpretation and a practical path through the current pipeline."),
        ("Narrative", "VF repertoire -> feature framework -> longitudinal stability/change -> cautious clinical-status overlay -> handover."),
        ("Primary guardrail", "Do not present this as a recurrent-UTI-only deck; UTI status is a later exploratory annotation."),
    ]):
        set_cell_fill(meta.cell(i, 0), LIGHT_GREY)
        p = meta.cell(i, 0).paragraphs[0]
        add_run(p, k, bold=True, color=BLUE)
        p2 = meta.cell(i, 1).paragraphs[0]
        add_run(p2, v)

    doc.add_paragraph()
    add_callout(
        doc,
        "Avoid saying",
        "Do not present legacy ASB-versus-UTI conclusions, older lpf significance claims, project-wide AMR association claims, or language that stable VF profiles establish a host-state mechanism.",
        fill=LIGHT_ORANGE,
        title_color=ORANGE,
    )

    doc.add_heading("Spoken Slides", level=1)
    for s in slides:
        doc.add_heading(f"Slide {s['n']}: {s['title']}", level=2)
        add_label_para(doc, "Target time", s["time"])
        add_label_para(doc, "Primary message", s["message"])
        p = doc.add_paragraph()
        add_run(p, "Script: ", bold=True, color=BLUE)
        add_run(p, s["script"])
        add_label_para(doc, "Transition", s["transition"])
        add_label_para(doc, "Caution", s["caution"])

    doc.add_page_break()
    doc.add_heading("Appendix Prompts", level=1)
    intro = doc.add_paragraph()
    add_run(intro, "Use these only when the discussion needs backup evidence, operational handover, or provenance.", color=MUTED)

    for n, title_text, prompt in appendix:
        doc.add_heading(f"Slide {n}: {title_text}", level=2)
        add_bullet(doc, prompt)

    doc.add_heading("Rehearsal Checklist", level=1)
    for item in [
        "By slide 5, the listener should understand what the VF matrix, modules, and outputs are.",
        "By slide 8, the listener should be able to state the central longitudinal stability result.",
        "By slide 11, the listener should hear that clinical-status VF associations are exploratory and not FDR-significant.",
        "By slide 14 or 22, the colleague should know which canonical files to open for rerun or provenance questions.",
    ]:
        add_bullet(doc, item)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
