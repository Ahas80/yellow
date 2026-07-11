#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE = Path("/Users/Aamir/Desktop/rUTIs/outputs/manual-20260527-current-review/presentations/ruti-longitudinal-vf-pipeline-review-v5")
OUT = WORKSPACE / "output" / "Longitudinal_Urinary_Ecoli_VF_Pipeline_Presenter_Guide_Compact_Onboarding_2026-05-28.docx"
MD_OUT = WORKSPACE / "output" / "Longitudinal_Urinary_Ecoli_VF_Pipeline_Presenter_Guide_Compact_Onboarding_2026-05-28.md"

BLUE = "2B6CB0"
GREEN = "2F855A"
ORANGE = "D97706"
RUST = "B65A3C"
MUTED = "64748B"
INK = "111827"
LIGHT = "F8FAFC"
CALLOUT = "EEF3F8"
CAUTION = "FFF4E6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.25


def add_para(doc, text: str, *, bold=False, color=INK, size=11, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)
    return p


def add_callout(doc, title: str, body: str, fill: str = CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)
    run.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(INK)
    return table


SLIDES = [
    {
        "n": 1,
        "title": "Virulence-factor profiling of longitudinal urinary E. coli isolates",
        "time": "1 min",
        "message": "Open with a VF-first project frame: the aim is to understand the bacterial VF repertoire and its longitudinal stability, with clinical status used later as an annotation.",
        "script": "I would start by saying: this project is not only a recurrent UTI story. The core object is a longitudinal set of urinary E. coli isolates, each converted into a virulence-factor profile. The three questions are: what VF repertoire is present, how stable those profiles are across repeated isolates from the same resident, and how carefully we can overlay UTI versus Not_UTI status. The status comparison matters, but it enters after we understand the data product and the longitudinal bacterial signal.",
        "transition": "Before looking at plots, I want to give the plain-English loop of what the pipeline is doing.",
        "caution": "Do not let the title drift back into a recurrent-UTI-only framing.",
    },
    {
        "n": 2,
        "title": "The YELLOW routine in one plain-English loop",
        "time": "1.5 min",
        "message": "Give the newcomer a mental model for the entire routine without scripts or dense counts.",
        "script": "Talk through the slide left to right. A clinical episode is a urine-sampling event. From that episode, we have a urinary E. coli isolate. Sequencing turns that isolate into a genome assembly. The VF pipeline asks which virulence-factor genes are detected in that assembly, so one isolate becomes one VF gene row. Those rows can then be summarised into modules, burden, similarity, and gain or loss. Only after that do we overlay clinical status. The key phrase is: one episode becomes one bacterial profile, and repeated profiles are compared over time.",
        "transition": "The next slide explains why the denominator changes depending on which object we are counting.",
        "caution": "Do not mention script names here; this slide is intentionally non-technical.",
    },
    {
        "n": 3,
        "title": "Unit changes explain denominator changes",
        "time": "1.5 min",
        "message": "Make denominator changes feel logical rather than alarming.",
        "script": "This slide is useful for preventing confusion. The project begins with 583 included clinical episodes, but not every clinical episode has a selected VF-ready assembly. That is why the VF-ready dataset is 556 episodes. From those 556 rows, we represent 227 VF genes and 32 modules. When we move to longitudinal analysis, the unit changes again: now we are counting consecutive within-resident pairs, giving 394 comparisons. When we focus on Not_UTI to UTI clinical transitions, the unit narrows to 11 transition cases. So the counts are not contradictory; they answer different questions.",
        "transition": "Because UTI status is used later, I will define exactly how an episode receives that label.",
        "caution": "Avoid comparing the 583, 556, 394, and 11 denominators as if they are the same denominator.",
    },
    {
        "n": 4,
        "title": "How an episode becomes UTI in this project",
        "time": "1.5 min",
        "message": "Clarify that UTI is defined by culture support plus catheter-aware symptoms.",
        "script": "Here I would say: in this analysis, UTI is not simply bacteria in urine. A UTI episode requires culture support and compatible symptoms, and the symptom rule is catheter-aware. If both conditions are met, the episode is UTI. If not, it sits in Not_UTI. That means Not_UTI is deliberately heterogeneous: it can include bacteriuria without the symptom rule, episodes that do not meet culture support, or other near-miss contexts. This matters because later status comparisons are comparing a small UTI group against a broad comparator.",
        "transition": "Now that the clinical annotation is defined, we can return to the VF-ready data asset.",
        "caution": "Do not describe Not_UTI as healthy controls.",
    },
    {
        "n": 5,
        "title": "VF-ready data asset",
        "time": "1 min",
        "message": "Show what the dataset physically represents: repeated isolates with selected assemblies and VF profiles.",
        "script": "The working VF dataset has 556 episode-level selected assemblies from 162 participants. Each row is a selected assembly linked back to the clinical episode. Across those rows, the pipeline records 227 binary VF gene columns, then groups those genes into 32 curated modules, including 18 UPEC-candidate modules. The timeline schematic is deliberately simple: repeated isolates from the same resident can carry different status labels, but the bacterial profiles are what we compare longitudinally.",
        "transition": "The next slide explains how those VF profiles are encoded and summarised.",
        "caution": "Do not display or state a total clinical participant count outside the VF-ready dataset.",
    },
    {
        "n": 6,
        "title": "VF features: genes, modules, and similarity",
        "time": "1.5 min",
        "message": "Explain the feature representation and why modules help interpretation.",
        "script": "This is the translation from genome to analysis features. The simplest object is the binary gene matrix: present or absent for each VF gene in each isolate. Because individual genes can be hard to interpret, the pipeline also maps genes into curated biological modules. Those modules help us talk about adhesion, iron acquisition, secretion, capsule and surface structures, toxins, and unassigned material. Then the longitudinal outputs compare profiles across time using burden, prevalence, Jaccard similarity, and gain/loss summaries.",
        "transition": "With that feature language in place, we can look at the observed VF repertoire.",
        "caution": "Modules are curation units, not validated disease prediction scores.",
    },
    {
        "n": 7,
        "title": "VF repertoire summarised as genes and modules",
        "time": "1.5 min",
        "message": "Separate descriptive prevalence from association testing.",
        "script": "Use the left panel as a descriptive prevalence ranking: these are common VF genes among the 556 VF-ready isolates. That does not mean they cause symptoms or distinguish UTI. The right panel shows how genes are distributed across curated modules. I would emphasise that this gives the colleague vocabulary for later slides. We are first learning what is present and how it is organised, before asking whether any clinical annotation appears to line up with it.",
        "transition": "The central result comes next: how similar repeated profiles are within the same resident.",
        "caution": "Do not interpret high prevalence as clinical importance by itself.",
    },
    {
        "n": 8,
        "title": "Most repeated within-resident VF profiles are highly stable",
        "time": "2 min",
        "message": "Present the main longitudinal VF result.",
        "script": "This is one of the main biological messages. The analysis considers 394 consecutive within-resident comparisons from 144 participants. The median Jaccard similarity is 1.000, and 62.4% of consecutive comparisons show no VF change. In plain terms, many repeated urinary E. coli isolates from the same resident have extremely similar VF presence/absence profiles. That suggests that, for many residents, the measured VF repertoire is conserved over time.",
        "transition": "The next question is what we do with the comparisons that are not perfectly stable.",
        "caution": "High VF similarity supports stability, but does not prove same-strain persistence without SNP or lineage context.",
    },
    {
        "n": 9,
        "title": "Gain/loss summaries flag candidates for follow-up",
        "time": "1.5 min",
        "message": "Explain how observed VF changes should be interpreted cautiously.",
        "script": "This plot looks at where VF genes appear to be gained or lost across consecutive pairs. I would frame this as a triage view, not a final mechanism. A gain or loss can happen because the resident has a replacement isolate, because of assembly or calling variation, or because there is genuine gene-content change. The value of the slide is that it tells us where to look more closely, especially when combined with ST and SNP context.",
        "transition": "That is why we next bring in sequence-type consistency.",
        "caution": "Avoid saying that a plotted gain or loss automatically represents biological evolution in the resident.",
    },
    {
        "n": 10,
        "title": "Sequence-type consistency helps interpret VF stability",
        "time": "1.5 min",
        "message": "Use lineage as context for interpreting profile stability or change.",
        "script": "Here the point is that VF profile stability needs lineage context. If repeated isolates share the same sequence type and have very similar VF profiles, that supports a persistent-lineage interpretation. If sequence type changes, then a VF change may be more consistent with replacement. But ST is still a coarse diagnostic. It helps organise the evidence; it does not independently prove same strain.",
        "transition": "Only after this VF and lineage structure do we move to the clinical-status overlay.",
        "caution": "Do not describe ST agreement as definitive persistence.",
    },
    {
        "n": 11,
        "title": "Clinical-status VF signals remain exploratory",
        "time": "2 min",
        "message": "Make the status analysis useful but bounded.",
        "script": "This is where UTI versus Not_UTI enters the analysis as a clinical annotation. The VF/model-ready dataset has 556 episodes, but only 17 are UTI and 539 are Not_UTI. The three blocks show the intended logic more clearly than the dense model plot: first, a simple screen can produce nominal-looking VF differences; second, after participant-aware modelling and FDR correction, there are 0 robust global VF hits; third, sparse/separation flags and lineage structure mean we should treat this as a prioritisation layer. The safe conclusion is that the current data support hypothesis generation, not a confirmed UTI-associated VF signature.",
        "transition": "The next slide shows why the longitudinal case view is still biologically useful.",
        "caution": "Do not mention old significant lpf claims or imply a corrected global association.",
    },
    {
        "n": 12,
        "title": "Participant 20026 transition example",
        "time": "2 min",
        "message": "Make the longitudinal argument concrete.",
        "script": "This is a worked example. Participant 20026 moves from Not_UTI to UTI over 42 days. The isolates are separated by only 5 SNPs and show no VF gene gain or loss in the measured profile. That means the measured virulence repertoire appears stable while the clinical state changes. The careful interpretation is that stable VF plus low SNP distance supports a host-state, timing, or unmeasured-regulation hypothesis. It does not prove the mechanism. Across the broader casebook there are 11 clinical Not_UTI to UTI transitions, 10 linked to WGS/VF evidence, with 4 stable-profile and 3 replacement-consistent transitions.",
        "transition": "I would close the spoken deck by turning those observations into takeaways and open decisions.",
        "caution": "Never say stable VF proves host mechanism.",
    },
    {
        "n": 13,
        "title": "Takeaways, limitations, and next steps",
        "time": "1 min",
        "message": "Close the spoken narrative with three clear conclusions.",
        "script": "The wrap-up is: first, the current pipeline produces a usable longitudinal VF dataset. Second, VF profiles are commonly stable within residents across repeated urinary E. coli isolates. Third, clinical-status associations remain exploratory because the UTI denominator is sparse and repeated measures and lineage structure matter. The next research step is not to overclaim a single VF marker, but to use this pipeline to prioritise lineage-aware follow-up, expression or regulation hypotheses, and careful review of transition cases.",
        "transition": "Then pause for questions, using the appendix slides depending on what the colleague asks.",
        "caution": "Do not turn the limitations into an apology; they define the correct next analyses.",
    },
    {
        "n": 14,
        "title": "Appendix: practical handover map",
        "time": "Q&A",
        "message": "Use when the colleague asks where to start or how to rerun the project.",
        "script": "Point them first to the canonical outputs: status_map.csv for clinical annotation, vf_pa_all.csv for the raw VF matrix, vf_analysis_ready.csv for the 556-row analysis dataset, and the gene/module map for curation. For reruns, walk from clinical status through assembly selection, VF matrix building, analysis-ready construction, modules and scores, longitudinal summaries, then the final figure pack.",
        "transition": "Use this slide as the practical handover rather than trying to explain the repo from memory.",
        "caution": "Do not ask a new colleague to begin with every script; begin with outputs and then trace backward.",
    },
    {
        "n": 15,
        "title": "Appendix: detailed clinical-to-VF pipeline",
        "time": "Q&A",
        "message": "Use for behind-the-scenes denominator and script-order questions.",
        "script": "This is the dense map. It shows why the simplified count ladder is true: 585 classified episodes before exclusions, 583 primary included clinical episodes, 1,291 assembly-level QC records including assembler alternatives, 556 selected assemblies and VF rows, then the 556-row VF/model-ready dataset. It also shows the feature framework and the transition application. This is the slide to use when someone asks exactly where rows are lost or why assembly records exceed episode records.",
        "transition": "If they ask about transition mechanisms specifically, move to the casebook slide.",
        "caution": "Keep the unit-of-analysis language explicit.",
    },
    {
        "n": 16,
        "title": "Appendix: full transition mechanism casebook",
        "time": "Q&A",
        "message": "Use for detailed questions about Not_UTI to UTI transitions.",
        "script": "This appendix slide shows the broader transition casebook. The safe statement is that 11 clinical Not_UTI to UTI transitions were identified; 10 have WGS/VF-linked evidence and 1 is missing the endpoint. Four transitions are stable-profile cases and three are replacement-consistent. These are evidence buckets, not proven mechanisms.",
        "transition": "For population-level signal or confounding questions, go to the final appendix slide.",
        "caution": "Do not make the transition buckets sound mutually exhaustive or mechanistically proven.",
    },
    {
        "n": 17,
        "title": "Appendix: robustness and lineage diagnostics",
        "time": "Q&A",
        "message": "Use when someone asks whether the status signal is robust or lineage-confounded.",
        "script": "The left panel reinforces the population-level boundary: no global VF association is significant after FDR correction. The right panel is a lineage diagnostic, showing why sequence type and bacterial population structure must be considered before interpreting status differences. This slide is there to keep the discussion honest if someone asks whether there is a global VF signature of UTI in the current dataset.",
        "transition": "Return to the main takeaway: strong descriptive VF pipeline, cautious clinical overlay.",
        "caution": "Do not introduce a global VF-plus-AMR claim; AMR is not part of this compact spoken deck.",
    },
]


def add_slide_section(doc: Document, item: dict) -> None:
    doc.add_heading(f"Slide {item['n']}. {item['title']}", level=2)
    meta = doc.add_paragraph()
    meta.add_run("Target time: ").bold = True
    meta.add_run(item["time"])
    meta.add_run("  |  ")
    meta.add_run("Primary message: ").bold = True
    meta.add_run(item["message"])
    add_para(doc, "How to say it live", bold=True, color=BLUE, size=11)
    add_para(doc, item["script"])
    add_para(doc, "Transition", bold=True, color=GREEN, size=11)
    add_para(doc, item["transition"])
    add_callout(doc, "Caution / avoid overstating", item["caution"], fill=CAUTION)


def build_markdown() -> str:
    lines = [
        "# Longitudinal urinary E. coli VF pipeline review: compact onboarding presenter guide",
        "",
        "**Purpose:** Help introduce a new team member to the current VF-first longitudinal pipeline and its cautious clinical annotation layer.",
        "",
        "**Timing:** 20-22 min presentation plus discussion.",
        "",
        "## Avoid Saying",
        "- Do not frame the current story as an ASB-versus-UTI analysis.",
        "- Do not claim a significant lpf result as current.",
        "- Do not claim a global AMR association.",
        "- Do not say stable VF proves host mechanism.",
        "",
    ]
    for s in SLIDES:
        lines.extend([
            f"## Slide {s['n']}. {s['title']}",
            f"**Target time:** {s['time']}",
            "",
            f"**Primary message:** {s['message']}",
            "",
            "**How to say it live:**",
            s["script"],
            "",
            f"**Transition:** {s['transition']}",
            "",
            f"**Caution:** {s['caution']}",
            "",
        ])
    return "\n".join(lines)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Longitudinal urinary E. coli VF pipeline review")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Compact onboarding presenter guide, aligned to the 17-slide v5 deck").italic = True
    subtitle.paragraph_format.space_after = Pt(12)

    add_callout(
        doc,
        "Purpose",
        "Use this guide to talk a new colleague through what the pipeline is doing, why denominators change, what the VF outputs mean, and where clinical status can be overlaid cautiously.",
        fill=CALLOUT,
    )

    doc.add_heading("Avoid saying", level=1)
    for item in [
        "Do not frame the current story as an ASB-versus-UTI analysis.",
        "Do not claim a significant lpf result as current.",
        "Do not claim a global AMR association.",
        "Do not say stable VF proves host mechanism.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Slide-by-slide script", level=1)
    for slide in SLIDES:
        add_slide_section(doc, slide)

    doc.add_heading("Removed from the compact spoken deck", level=1)
    add_para(
        doc,
        "The compact v5 deck removes redundant status-detail, module-by-status, AMR backup, and duplicate stable-strain/context slides. Those materials can remain in the fuller v4 deck for deep-dive backup, but they are intentionally absent from the live compact version.",
    )

    doc.save(OUT)
    MD_OUT.write_text(build_markdown(), encoding="utf-8")
    print(OUT)
    print(MD_OUT)


if __name__ == "__main__":
    main()
