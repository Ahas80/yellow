#!/usr/bin/env python3

"""Build registry-bound Longcycler-only presenter guides in DOCX and Markdown."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
from pathlib import Path
import tempfile
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path("/Users/Aamir/Desktop/rUTIs")
CANONICAL_PRESENTATION_ROOT = PROJECT_ROOT / "outputs" / "manual-20260527-current-review" / "presentations"
REGISTRY_DEFAULT = PROJECT_ROOT / "results" / "pipeline" / "longcycler_release_claim_registry.json"
FORBIDDEN = "".join(("fl", "ye"))

BLUE = "2B6CB0"
GREEN = "2F855A"
ORANGE = "D97706"
INK = "111827"
MUTED = "64748B"
LIGHT = "F8FAFC"
LINE = "CBD5E1"


def development_registry() -> dict:
    return {
        "schema_version": "longcycler_release_claim_registry_v1",
        "generated_at": "2026-07-13 12:00:00 CEST",
        "analysis_scope": {
            "assembly_policy": "selected QC-passing Longcycler only",
            "clinical_phenotype": "operational UTI phenotype",
            "clinical_definition_version": "operational-2026-07",
            "interpretation": "exploratory observational analysis; no causal claim",
        },
        "method_contract": {
            "operational_phenotype": {
                "culture_lower_bound_cfu_per_ml": 1000,
                "rule": "versioned operational culture-plus-compatible-symptom phenotype",
                "caveat": "not a reconstruction of the full published protocol",
            },
            "assembly_qc": {
                "max_contigs": 200,
                "min_n50_bp": 20000,
                "min_genome_size_bp": 4000000,
                "max_genome_size_bp": 6000000,
                "excluded_metrics": ["read coverage", "completeness", "contamination"],
            },
            "vfdb": {
                "tool": "ABRicate",
                "database": "VFDB",
                "min_identity_pct": 80,
                "min_coverage_pct": 80,
                "provenance": "SHA-bound calls from the selected Longcycler FASTA manifest",
            },
            "mlst": {
                "role": "lineage context; not pair-specific continuity proof",
                "provider_min_good_targets_pct": 95,
                "provider_policy": "provider_qc95 call key/path-linked to the selected Longcycler episode; local fallback excluded",
                "fallback": "labelled local MLST from the same selected Longcycler FASTA where required",
            },
            "direct_pair_evidence": {
                "tool": "dnadiff",
                "role": "primary pair-specific distance evidence",
                "operational_snp_threshold": 25,
                "priority": "graph connectivity and MLST agreement cannot override a conflicting direct pair",
            },
            "population_context": {
                "core_genome_tool": "Parsnp",
                "pangenome_tool": "Panaroo",
                "role": "population context; not a substitute for direct pair evidence",
            },
        },
        "analytical_cohort": {"episodes": 532, "residents": 161, "operational_UTI": 16, "operational_Not_UTI": 516},
        "attrition_qc_context": {"label": "full clinical source retained only for attrition/QC context", "episodes": 583, "residents": 166, "operational_UTI": 18, "operational_Not_UTI": 565},
        "direct_pairs": {"all_within_resident": 893},
        "adjacent_transitions": {"pairs": 371, "residents": 139, "operational_snp_threshold": 25, "at_or_below_threshold": 140, "Not_UTI_to_UTI": 9, "Not_UTI_to_UTI_at_or_below_threshold": 5},
        "mechanism_casebook": {"cases": 9, "linked": 9, "missing": 0},
        "near_miss_audit": {"rows": 17, "label": "near-miss rows; not operational UTI cases"},
        "selected_uti_event_genomes": {"genomes": 32, "residents": 29, "operational_UTI": 15, "operational_Not_UTI": 17},
        "genomic_dimensions": {"VFDB_binary_features": 227, "MLST_typed_episodes": 514, "distinct_preferred_ST_labels": 80},
        "research_questions": {"first": "RQ01", "last": "RQ10", "count": 10, "retired_questions": 0},
        "plot_files": [],
        "sources": [],
    }


def validate_registry(registry: dict) -> None:
    a = registry.get("analytical_cohort", {})
    x = registry.get("attrition_qc_context", {})
    t = registry.get("adjacent_transitions", {})
    c = registry.get("mechanism_casebook", {})
    r = registry.get("research_questions", {})
    scope = registry.get("analysis_scope", {})
    methods = registry.get("method_contract", {})
    phenotype = methods.get("operational_phenotype", {})
    assembly_qc = methods.get("assembly_qc", {})
    vfdb = methods.get("vfdb", {})
    mlst = methods.get("mlst", {})
    direct = methods.get("direct_pair_evidence", {})
    population = methods.get("population_context", {})
    checks = {
        "schema": registry.get("schema_version") == "longcycler_release_claim_registry_v1",
        "scope": (
            scope.get("assembly_policy"),
            scope.get("clinical_phenotype"),
            scope.get("interpretation"),
        ) == (
            "selected QC-passing Longcycler only",
            "operational UTI phenotype",
            "exploratory observational analysis; no causal claim",
        ),
        "cohort": (a.get("episodes"), a.get("residents"), a.get("operational_UTI"), a.get("operational_Not_UTI")) == (532, 161, 16, 516),
        "attrition/QC context": (
            x.get("label"), x.get("episodes"), x.get("residents"),
            x.get("operational_UTI"), x.get("operational_Not_UTI"),
        ) == ("full clinical source retained only for attrition/QC context", 583, 166, 18, 565),
        "direct pairs": registry.get("direct_pairs", {}).get("all_within_resident") == 893,
        "adjacent transitions": (t.get("pairs"), t.get("residents"), t.get("operational_snp_threshold"), t.get("at_or_below_threshold")) == (371, 139, 25, 140),
        "focused transitions": (t.get("Not_UTI_to_UTI"), t.get("Not_UTI_to_UTI_at_or_below_threshold")) == (9, 5),
        "casebook": (c.get("cases"), c.get("linked"), c.get("missing")) == (9, 9, 0),
        "near miss": registry.get("near_miss_audit", {}).get("rows") == 17,
        "research questions": (r.get("first"), r.get("last"), r.get("count"), r.get("retired_questions")) == ("RQ01", "RQ10", 10, 0),
        "operational phenotype method": (
            phenotype.get("culture_lower_bound_cfu_per_ml"), phenotype.get("rule"), phenotype.get("caveat"),
        ) == (1000, "versioned operational culture-plus-compatible-symptom phenotype", "not a reconstruction of the full published protocol"),
        "assembly QC method": (
            assembly_qc.get("max_contigs"), assembly_qc.get("min_n50_bp"),
            assembly_qc.get("min_genome_size_bp"), assembly_qc.get("max_genome_size_bp"),
            set(assembly_qc.get("excluded_metrics", [])),
        ) == (200, 20000, 4000000, 6000000, {"read coverage", "completeness", "contamination"}),
        "VFDB method": (
            vfdb.get("tool"), vfdb.get("database"), vfdb.get("min_identity_pct"),
            vfdb.get("min_coverage_pct"), vfdb.get("provenance"),
        ) == ("ABRicate", "VFDB", 80, 80, "SHA-bound calls from the selected Longcycler FASTA manifest"),
        "MLST method": (
            mlst.get("role"), mlst.get("provider_min_good_targets_pct"),
            mlst.get("provider_policy"), mlst.get("fallback"),
        ) == (
            "lineage context; not pair-specific continuity proof", 95,
            "provider_qc95 call key/path-linked to the selected Longcycler episode; local fallback excluded",
            "labelled local MLST from the same selected Longcycler FASTA where required",
        ),
        "direct-pair method": (
            direct.get("tool"), direct.get("role"), direct.get("operational_snp_threshold"), direct.get("priority"),
        ) == (
            "dnadiff", "primary pair-specific distance evidence", 25,
            "graph connectivity and MLST agreement cannot override a conflicting direct pair",
        ),
        "population-context methods": (
            population.get("core_genome_tool"), population.get("pangenome_tool"), population.get("role"),
        ) == ("Parsnp", "Panaroo", "population context; not a substitute for direct pair evidence"),
    }
    failures = [name for name, passed in checks.items() if not passed]
    if FORBIDDEN in json.dumps(registry).lower():
        failures.append("retired input token")
    if failures:
        raise ValueError("Claim registry contract failed: " + ", ".join(failures))


def counts(registry: dict) -> dict:
    a = registry["analytical_cohort"]
    x = registry["attrition_qc_context"]
    t = registry["adjacent_transitions"]
    c = registry["mechanism_casebook"]
    g = registry["genomic_dimensions"]
    return {
        "episodes": a["episodes"], "residents": a["residents"], "uti": a["operational_UTI"], "not_uti": a["operational_Not_UTI"],
        "source_episodes": x["episodes"], "source_residents": x["residents"], "source_uti": x["operational_UTI"], "source_not_uti": x["operational_Not_UTI"],
        "direct": registry["direct_pairs"]["all_within_resident"], "transitions": t["pairs"], "transition_residents": t["residents"],
        "threshold": t["operational_snp_threshold"], "at_threshold": t["at_or_below_threshold"], "to_uti": t["Not_UTI_to_UTI"], "to_uti_threshold": t["Not_UTI_to_UTI_at_or_below_threshold"],
        "cases": c["cases"], "linked": c["linked"], "missing": c["missing"], "near_miss": registry["near_miss_audit"]["rows"],
        "vf": g["VFDB_binary_features"], "mlst": g["MLST_typed_episodes"], "st": g["distinct_preferred_ST_labels"],
    }


def item(title: str, say: str, emphasise: str, boundary: str) -> dict:
    return {"title": title, "say": say, "emphasise": emphasise, "boundary": boundary}


def v3_items(c: dict) -> list[dict]:
    return [
        item("Title and release anchors", f"Open with the selected analytical denominator: {c['episodes']} episodes from {c['residents']} residents, with {c['transitions']} adjacent transitions and {c['uti']} operational UTI episodes.", "The whole review is selected QC-passing Longcycler only.", "Operational UTI is an exploratory annotation, not the organising genome source."),
        item("Scope lock", f"One selected assembly per episode creates {c['episodes']} episode-level genome profiles. The temporal layer is {c['transitions']} adjacent transitions from {c['transition_residents']} residents.", "Name the assembly policy and analysis unit before discussing results.", "Do not mix source/QC genomes into the analytical cohort."),
        item("Selected cohort", f"The release cohort is {c['episodes']} episodes, {c['residents']} residents, {c['uti']} operational UTI, and {c['not_uti']} operational Not_UTI.", f"There are {c['vf']} binary VF features in the current release.", "Not_UTI is a mixed comparator, not one biological state."),
        item("Numbered pipeline", "Walk through the validated clinical-to-VF handoffs and point to the RQ01–RQ10 release layer.", "Each stage consumes the same selected cohort.", "Do not introduce an alternate assembly path."),
        item("VF representation", f"Each selected episode is represented by {c['vf']} binary VF features, then summarised through curated modules and longitudinal comparisons.", "Genes remain the underlying evidence even when modules are shown.", "Presence/absence is not expression, activity, or causality."),
        item("Gene repertoire", "Read the ranked genes as a descriptive view of what is commonly detected in the selected cohort.", "The denominator is the selected episode-level genome set.", "Do not interpret prevalence ranking as a UTI association."),
        item("Curated modules", "Explain that modules organise genes into interpretable systems for navigation and summary.", "Curation improves readability without changing binary gene evidence.", "Modules are not validated disease-causality scores."),
        item("Pair layers", f"Separate {c['direct']} direct within-resident pairs from {c['transitions']} adjacent transitions. The first is broad; the second is temporally ordered.", f"{c['at_threshold']} adjacent transitions are at or below the operational {c['threshold']}-SNP threshold.", "Always name the pair unit beside the count."),
        item("Gain and loss", f"Use the {c['transitions']} adjacent transitions to identify candidate VF gains and losses for follow-up.", "Interpret change with lineage, genome distance, and QC context.", "A detected change may reflect replacement or calling differences as well as true gene-content change."),
        item("MLST context", f"Preferred MLST labels are available for {c['mlst']} selected episodes across {c['st']} distinct labels.", "Use lineage as a diagnostic layer before status interpretation.", "ST agreement does not prove the same strain."),
        item("Exploratory clinical overlay", f"The selected phenotype split is {c['uti']} operational UTI versus {c['not_uti']} operational Not_UTI episodes.", "The clinical model is hypothesis-generating and participant-aware.", "No causal or confirmatory VF status claim is supported by the sparse phenotype denominator."),
        item("Aggregate focused transitions", f"There are {c['to_uti']} focused Not_UTI-to-UTI adjacent transitions; {c['to_uti_threshold']} are at or below {c['threshold']} SNPs. The casebook is {c['cases']}/{c['linked']}/{c['missing']} cases/linked/missing.", "Report the aggregate casebook, not a single anecdote.", "Low genome distance cannot identify host-state change or bacterial regulation by itself."),
        item("Evidence, boundary, next steps", "Land on three messages: the selected longitudinal evidence is coherent; the clinical phenotype remains sparse; and follow-up must be lineage-aware.", "Ask what extra evidence would distinguish bacterial change from host-state and sampling effects.", "Keep causal interpretation outside the supported claim set."),
        item("Operational handover", "Point to the selected manifest, clinical key, VF matrix, pair tables, transition casebook, and RQ01–RQ10 runner.", "The release claim registry is the provenance anchor.", "Do not use an unregistered intermediate as a release source."),
        item("Denominator audit", f"The source clinical context is {c['source_episodes']} episodes and {c['source_residents']} residents ({c['source_uti']} UTI, {c['source_not_uti']} Not_UTI). It is retained only for attrition/QC context.", f"Analytical claims remain {c['episodes']}/{c['residents']}/{c['uti']}/{c['not_uti']}.", "Never substitute source/QC counts into analytical claims."),
        item("Module prevalence annotation", "Use module prevalence by status as a descriptive clinical annotation view.", "Keep the operational phenotype denominator visible.", "Repeated measures and lineage are not fully resolved by a simple prevalence plot."),
        item("Complete mechanism casebook", f"The focused mechanism casebook contains {c['cases']} cases, all {c['linked']} linked, with {c['missing']} missing endpoints.", "Use the evidence matrix to show heterogeneity across cases.", "Buckets organise evidence; they do not prove mechanism."),
        item("Stable strain and changing clinical state", f"Among the {c['to_uti']} focused transitions, {c['to_uti_threshold']} are at or below {c['threshold']} SNPs.", "Stable genomic context can motivate host-state or regulation hypotheses.", "It cannot establish either hypothesis without additional evidence."),
        item("Robustness boundary", f"The operational phenotype contains only {c['uti']} UTI episodes, and the {c['near_miss']}-row near-miss audit is separate.", "Use robustness diagnostics to define uncertainty.", "Near-miss rows are not operational UTI cases."),
        item("Lineage diagnostic", "Use the VF-profile PCoA by preferred sequence type to assess lineage structure.", "Lineage structure should be considered before interpreting status patterns.", "Clustering is diagnostic context, not causal evidence."),
        item("Accessory context", "Keep accessory and mobile-element information as transition-level context when it helps interpret a specific pattern.", "The VF review remains the scientific centre.", "No dedicated AMR association or causal accessory claim is made."),
        item("Release registry and sources", "Close the appendix by showing the registry path, SHA-256, selected denominators, pair layers, focused casebook, near-miss audit, and RQ01–RQ10 marker.", "Use the registry for provenance questions.", "Do not answer release-count questions from memory or an earlier deck."),
    ]


def onboarding_items(c: dict) -> list[dict]:
    return [
        item("Plain-English loop", "Walk from clinical episode to selected genome, binary VF row, longitudinal comparison, and finally the clinical overlay.", "The genome source is selected QC-passing Longcycler only.", "Do not collapse episodes, residents, pairs, and transitions into one unit."),
        item("Denominator units", f"Explain why the counts change: {c['source_episodes']} source episodes retained only for attrition/QC context, {c['episodes']} selected analytical episodes, {c['direct']} direct pairs, and {c['transitions']} adjacent transitions.", "Name the unit whenever a number appears.", "Source/QC counts do not replace analytical counts."),
        item("Operational phenotype", f"Explain the project annotation that yields {c['uti']} operational UTI and {c['not_uti']} operational Not_UTI selected episodes.", "Not_UTI is a mixed comparator.", "The operational label is not a universal clinical diagnosis."),
    ]


def v5_items(c: dict) -> list[dict]:
    base = v3_items(c)
    return [
        base[0],
        *onboarding_items(c),
        base[2],
        base[4],
        item("Paired repertoire views", "Show gene-level prevalence and curated module structure together as two descriptive views of the same selected cohort.", f"The current release contains {c['vf']} binary VF features.", "Neither plot is a causal clinical model."),
        item("Pair-layer stability", f"Describe VF similarity using {c['direct']} direct pairs and {c['transitions']} adjacent transitions from {c['transition_residents']} residents.", "Use adjacent transitions for temporal questions.", "Do not reuse an earlier plot denominator."),
        base[8], base[9], base[10], base[11], base[12], base[13], base[3], base[16],
        item("Robustness and lineage diagnostics", f"Keep the robustness plot beside the lineage PCoA. The {c['near_miss']}-row near-miss audit remains separate from operational UTI cases.", "These are backup interpretation aids.", "Neither diagnostic converts the sparse phenotype analysis into a confirmatory claim."),
    ]


def guide_items(variant: str, c: dict) -> list[dict]:
    if variant == "v3":
        return v3_items(c)
    if variant == "v4":
        base = v3_items(c)
        return [base[0], *onboarding_items(c), *base[1:]]
    if variant in {"v5-full", "v5-compact"}:
        return v5_items(c)
    raise ValueError(f"Unknown variant: {variant}")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    title = doc.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_after = Pt(8)

    for style_name, size, color, before, after in (
        ("Heading 1", 17, BLUE, 12, 5),
        ("Heading 2", 12.5, INK, 7, 3),
        ("Heading 3", 10.5, GREEN, 5, 2),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display" if style_name != "Heading 3" else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_labeled_paragraph(doc: Document, label: str, value: str, color: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.space_after = Pt(3)
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor.from_string(color)
    paragraph.add_run(value)


def build_docx(variant: str, registry: dict, registry_path: str, registry_sha: str, items: list[dict], output: Path) -> None:
    c = counts(registry)
    label = {
        "v3": "22-slide scientific review",
        "v4": "25-slide review with onboarding",
        "v5-full": "17-slide full review",
        "v5-compact": "17-slide compact onboarding review",
    }[variant]
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "LONGCYCLER-ONLY PRESENTER GUIDE  |  OPERATIONAL UTI PHENOTYPE"
    header.style = doc.styles["Caption"]
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(7.5)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Selected analytical release  •  ")
    run.font.name = "Aptos"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(footer)

    eyebrow = doc.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(8)
    run = eyebrow.add_run("PRESENTER FIELD GUIDE  /  CURRENT LONGCYCLER RELEASE")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    title = doc.add_paragraph(style="Title")
    title.add_run("Longitudinal urinary E. coli\nvirulence-factor review")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle_run = subtitle.add_run(f"Canonical guide for the {label}")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor.from_string(MUTED)

    meta = doc.add_table(rows=2, cols=3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    meta_values = [
        ("Analytical cohort", f"{c['episodes']} episodes / {c['residents']} residents"),
        ("Operational phenotype", f"{c['uti']} UTI / {c['not_uti']} Not_UTI"),
        ("Longitudinal evidence", f"{c['direct']} direct / {c['transitions']} adjacent"),
        ("Focused transition view", f"{c['to_uti']} cases / {c['to_uti_threshold']} ≤{c['threshold']} SNPs"),
        ("Mechanism casebook", f"{c['cases']}/{c['linked']}/{c['missing']} cases/linked/missing"),
        ("Research questions", "RQ01–RQ10"),
    ]
    for index, cell in enumerate(meta._cells):
        cell.width = Inches(2.25)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT if index % 2 == 0 else "FFFFFF")
        set_cell_border(cell)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        label_run = paragraph.add_run(meta_values[index][0] + "\n")
        label_run.bold = True
        label_run.font.size = Pt(8.2)
        label_run.font.color.rgb = RGBColor.from_string(BLUE)
        value_run = paragraph.add_run(meta_values[index][1])
        value_run.font.size = Pt(8.8)
        value_run.font.color.rgb = RGBColor.from_string(INK)

    doc.add_heading("How to use this guide", level=1)
    add_labeled_paragraph(doc, "Say", "Use the short script as the spoken spine; pause on plots and let the evidence carry the slide.", BLUE)
    add_labeled_paragraph(doc, "Emphasise", "Every analytical claim is registry-bound and selected Longcycler only.", GREEN)
    add_labeled_paragraph(doc, "Boundary", "Operational UTI is exploratory observational context; no causal claim is supported.", ORANGE)

    doc.add_heading("Release provenance", level=2)
    provenance = doc.add_paragraph()
    provenance.paragraph_format.space_after = Pt(0)
    provenance.add_run("Registry: ").bold = True
    provenance.add_run(registry_path)
    provenance.add_run("\nSHA-256: ").bold = True
    provenance.add_run(registry_sha)

    doc.add_page_break()
    doc.add_heading("Slide-by-slide script", level=1)
    intro = doc.add_paragraph(f"{len(items)} slides. Suggested cadence: 35–70 seconds per spoken slide; appendix slides are question-led.")
    intro.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    for index, entry in enumerate(items, start=1):
        heading = doc.add_heading(f"Slide {index:02d} — {entry['title']}", level=2)
        heading.paragraph_format.keep_with_next = True
        add_labeled_paragraph(doc, "Say", entry["say"], BLUE)
        add_labeled_paragraph(doc, "Emphasise", entry["emphasise"], GREEN)
        add_labeled_paragraph(doc, "Boundary", entry["boundary"], ORANGE)
        if index in {4, 8, 12, 16, 20, 24} and index < len(items):
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    doc.add_page_break()
    doc.add_heading("Rapid Q&A guardrails", level=1)
    qa = [
        ("Which genome source?", "One selected QC-passing Longcycler assembly per analytical episode."),
        ("Which analytical denominator?", f"{c['episodes']} episodes from {c['residents']} residents; {c['uti']} operational UTI and {c['not_uti']} operational Not_UTI."),
        ("Why do pair counts differ?", f"{c['direct']} direct pairs include all within-resident comparisons; {c['transitions']} adjacent transitions are temporally ordered."),
        ("What does ≤25 SNPs mean?", f"It is the operational threshold used for context. {c['at_threshold']} adjacent transitions meet it; {c['to_uti_threshold']} of {c['to_uti']} focused transitions meet it."),
        ("Is the casebook complete?", f"Yes: {c['cases']} cases, {c['linked']} linked, {c['missing']} missing."),
        ("Are near-miss rows UTI cases?", f"No. The {c['near_miss']}-row audit is explicitly separate from operational UTI cases."),
        ("What is the claim level?", "Exploratory observational analysis. No causal claim."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    table.rows[0].cells[0].text = "Question"
    table.rows[0].cells[1].text = "Release-safe answer"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8.5)
    for question, answer in qa:
        row = table.add_row().cells
        row[0].text = question
        row[1].text = answer
        row[0].width = Inches(2.1)
        row[1].width = Inches(4.8)
        for cell in row:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1.5)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)

    doc.add_heading("Registered sources", level=1)
    sources = registry.get("sources", [])
    if sources:
        source_table = doc.add_table(rows=1, cols=3)
        source_table.style = "Table Grid"
        source_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ("Role", "Path", "SHA-256")
        for idx, value in enumerate(headers):
            source_table.rows[0].cells[idx].text = value
            set_cell_shading(source_table.rows[0].cells[idx], BLUE)
            for run in source_table.rows[0].cells[idx].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(7.5)
        set_repeat_table_header(source_table.rows[0])
        for source in sources:
            row = source_table.add_row().cells
            row[0].text = str(source.get("role", ""))
            row[1].text = str(source.get("path", ""))
            row[2].text = str(source.get("sha256", ""))
            for cell in row:
                set_cell_border(cell)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = "Aptos Mono" if cell != row[0] else "Aptos"
                        run.font.size = Pt(6.5)
    else:
        doc.add_paragraph("Development fixture: final registered sources will populate at canonical regeneration.")

    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(prefix=f".{output.stem}.", suffix=".tmp.docx", dir=output.parent, delete=False) as handle:
        temp_path = Path(handle.name)
    try:
        doc.save(temp_path)
        validate_docx(temp_path, len(items))
        os.replace(temp_path, output)
    finally:
        temp_path.unlink(missing_ok=True)


def build_markdown(variant: str, registry: dict, registry_path: str, registry_sha: str, items: list[dict]) -> str:
    c = counts(registry)
    lines = [
        "# Longitudinal urinary E. coli virulence-factor review",
        "",
        f"Canonical Longcycler-only presenter guide for `{variant}`.",
        "",
        "## Release anchors",
        "",
        f"- Analytical cohort: {c['episodes']} episodes; {c['residents']} residents; {c['uti']} operational UTI; {c['not_uti']} operational Not_UTI.",
        f"- Longitudinal evidence: {c['direct']} direct pairs; {c['transitions']} adjacent transitions from {c['transition_residents']} residents; {c['at_threshold']} at or below {c['threshold']} SNPs.",
        f"- Focused transition view: {c['to_uti']} Not_UTI→UTI; {c['to_uti_threshold']} at or below {c['threshold']} SNPs; casebook {c['cases']}/{c['linked']}/{c['missing']} cases/linked/missing.",
        f"- Near-miss audit: {c['near_miss']} rows, separate from operational UTI cases.",
        "- Research questions: RQ01–RQ10.",
        "- Interpretation: exploratory observational analysis; no causal claim.",
        "",
        "## Release provenance",
        "",
        f"- Registry: `{registry_path}`",
        f"- Registry SHA-256: `{registry_sha}`",
        "",
        "## Slide-by-slide script",
        "",
    ]
    for index, entry in enumerate(items, start=1):
        lines.extend([
            f"### Slide {index:02d} — {entry['title']}",
            "",
            f"**Say:** {entry['say']}",
            "",
            f"**Emphasise:** {entry['emphasise']}",
            "",
            f"**Boundary:** {entry['boundary']}",
            "",
        ])
    return "\n".join(lines).rstrip() + "\n"


def validate_docx(path: Path, expected_slides: int) -> None:
    if path.stat().st_size < 10_000:
        raise ValueError(f"DOCX is unexpectedly small: {path}")
    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore").lower()
        if FORBIDDEN in document_xml:
            raise ValueError("Retired input token remains in DOCX content")
        slide_headers = document_xml.count("slide ")
        if slide_headers < expected_slides:
            raise ValueError(f"Expected at least {expected_slides} slide headings, found {slide_headers}")


def atomic_write_text(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(prefix=f".{path.name}.", suffix=".tmp", dir=path.parent, delete=False, mode="w", encoding="utf-8") as handle:
        handle.write(value)
        temp_path = Path(handle.name)
    try:
        os.replace(temp_path, path)
    finally:
        temp_path.unlink(missing_ok=True)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--variant", required=True, choices=("v3", "v4", "v5-full", "v5-compact"))
    parser.add_argument("--registry", type=Path, default=REGISTRY_DEFAULT)
    parser.add_argument("--output-docx", type=Path, required=True)
    parser.add_argument("--output-md", type=Path, required=True)
    parser.add_argument("--dev-fixture", action="store_true")
    args = parser.parse_args()

    output_docx = args.output_docx.resolve()
    output_md = args.output_md.resolve()
    if args.dev_fixture:
        for output in (output_docx, output_md):
            if str(output).startswith(str(CANONICAL_PRESENTATION_ROOT.resolve()) + os.sep):
                raise ValueError("Development fixtures may only write outside the canonical presentation tree")
        registry = development_registry()
        registry_bytes = (json.dumps(registry, sort_keys=True, indent=2) + "\n").encode("utf-8")
        registry_path = "development fixture generated in external scratch"
    else:
        registry_bytes = args.registry.read_bytes()
        registry = json.loads(registry_bytes)
        registry_path = str(args.registry.resolve())
    validate_registry(registry)
    registry_sha = hashlib.sha256(registry_bytes).hexdigest()
    items = guide_items(args.variant, counts(registry))
    expected = {"v3": 22, "v4": 25, "v5-full": 17, "v5-compact": 17}[args.variant]
    if len(items) != expected:
        raise ValueError(f"{args.variant}: expected {expected} guide entries, found {len(items)}")
    markdown = build_markdown(args.variant, registry, registry_path, registry_sha, items)
    if FORBIDDEN in markdown.lower():
        raise ValueError("Retired input token remains in Markdown guide")
    build_docx(args.variant, registry, registry_path, registry_sha, items, output_docx)
    atomic_write_text(output_md, markdown)
    if not args.dev_fixture and hashlib.sha256(args.registry.read_bytes()).hexdigest() != registry_sha:
        raise ValueError("Claim-registry bytes changed while the presenter guide was generated")
    print(json.dumps({"variant": args.variant, "slides": len(items), "docx": str(output_docx), "markdown": str(output_md), "registry": registry_path, "registry_sha256": registry_sha}, indent=2))


if __name__ == "__main__":
    main()
