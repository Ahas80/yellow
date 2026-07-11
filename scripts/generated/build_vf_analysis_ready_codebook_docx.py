#!/usr/bin/env python3
"""
Build a lay-friendly Word codebook for results/vf/vf_analysis_ready.csv.

The document is generated from the live dataset and companion mapping files so
the appendix stays synchronized with the actual headers.
"""

from __future__ import annotations

import math
from collections import Counter
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/Aamir/Desktop/rUTIs")
VF_READY = ROOT / "results/vf/vf_analysis_ready.csv"
VF_PA = ROOT / "results/vf/vf_pa_all.csv"
GENE_MAP = ROOT / "results/vf/gene_map.csv"
DIAGNOSTICS = ROOT / "results/vf/vf_dataset_diagnostics.txt"
OUT_DOCX = ROOT / "outputs/codebooks/vf_analysis_ready_lay_codebook.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "4B5563"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_YELLOW = "FFF7D6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in: float) -> None:
    width = int(width_in * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_width(table, widths: list[float], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_dxa = int(sum(widths) * 1440)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, widths: list[float], header_fill: str = LIGHT_BLUE, font_size: int = 8) -> None:
    set_table_width(table, widths)
    for row_i, row in enumerate(table.rows):
        set_row_cant_split(row)
        if row_i == 0:
            set_repeat_table_header(row)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor.from_string(INK)
            if row_i == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], font_size: int = 8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_row = table.rows[0]
    for idx, text in enumerate(headers):
        header_row.cells[idx].text = text
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = "" if pd.isna(text) else str(text)
    style_table(table, widths, font_size=font_size)
    return table


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_note_box(doc: Document, title: str, text: str, fill: str = PALE_YELLOW):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.add_run(text)
    set_table_width(table, [6.5])
    return table


def clean_value(x) -> str:
    if pd.isna(x):
        return "blank"
    if isinstance(x, float) and x.is_integer():
        return str(int(x))
    return str(x)


def diagnostic_value(diag_text: str, prefix: str) -> str:
    for line in diag_text.splitlines():
        if line.startswith(prefix):
            return line.split(":", 1)[1].strip()
    return "not reported"


def summarize_values(series: pd.Series, is_gene: bool = False) -> str:
    if is_gene:
        return "0 = gene not detected; 1 = gene detected."
    non_missing = series.dropna()
    if non_missing.empty:
        return "All values are blank in the current file."
    if pd.api.types.is_bool_dtype(series):
        return "TRUE = yes/condition met; FALSE = no/condition not met; blank = not recorded or not applicable."
    unique = sorted({clean_value(v) for v in non_missing.unique()})
    if len(unique) <= 8:
        return "Values seen: " + ", ".join(unique) + "."
    if pd.api.types.is_numeric_dtype(series):
        vals = pd.to_numeric(non_missing, errors="coerce").dropna()
        if vals.empty:
            return "Mixed values; see source file."
        return (
            f"Numeric values. Current range: {clean_value(vals.min())} to "
            f"{clean_value(vals.max())}; median: {clean_value(vals.median())}."
        )
    counts = Counter(clean_value(v) for v in non_missing)
    common = ", ".join(f"{k} ({v})" for k, v in counts.most_common(5))
    return f"Text values. Most common in current file: {common}."


def missing_text(series: pd.Series, column: str, group: str) -> str:
    n_missing = int(series.isna().sum())
    if n_missing == 0:
        return "No blanks in the current file."
    if column in {"ST", "ST_provider", "provider_PercGoodTargets", "provider_file", "provider_batch_match", "provider_assembler"}:
        return f"{n_missing} blank(s): no usable provider ST detail was available for those rows."
    if column == "Not_UTI_subgroup":
        return f"{n_missing} blank(s): this subgroup is blank for rows classified as UTI."
    if group == "QC/curation":
        return f"{n_missing} blank(s): usually not applicable because no exclusion, duplicate, or manual note was needed."
    return f"{n_missing} blank(s): not recorded, not applicable, or not available in the current file."


def group_for_column(col: str, gene_cols: set[str]) -> str:
    if col in gene_cols:
        return "VF gene"
    if col in {"Participant_id", "tp_lab", "Episode_ID", "Event_type", "Collection_Date", "Batch"}:
        return "Identifier"
    if col in {"ST", "ST_source", "ST_provider", "ST_local", "provider_PercGoodTargets", "provider_file", "provider_batch_match", "provider_assembler"}:
        return "ST/MLST"
    if col in {
        "UTI_Status", "Primary_Status", "UTI_binary", "Not_UTI_subgroup", "Infection_Status",
        "Infection_Status_legacy", "Infection_Status_old", "UTI_definition_version",
        "UTI_classification_confidence", "UTI_classification_reason", "Status_Confidence_epi",
        "Sx_source_epi", "UTI_Label"
    }:
        return "Clinical status"
    if col.startswith("cat_") or col in {
        "vf_count_total", "total_vf_count_all", "total_vf_count_curated",
        "total_vf_count_upec_candidate", "total_vf_count_unassigned",
        "low_confidence_count", "n_timepoints"
    }:
        return "VF summary"
    if col in {
        "Urine_collection_method", "urine_collection_method_raw", "urine_collection_method_norm",
        "catheter_rule", "symptom_compatible_uti", "symptom_rule_met",
        "local_urinary_symptom_any", "systemic_symptom_any", "flankpain_present",
        "dysuria_present", "urgency_present", "frequency_present", "incontinence_present",
        "pus_present", "fever_present", "rigors_present", "delirium_present",
        "suprapubic_pain_present", "other_sxs_present", "cfu_raw", "cfu_raw_parsed",
        "cfu_ge_1e3", "cfu_ge_1e4", "cfu_ge_1e5", "culture_supports_uti",
        "cfu_threshold_used_for_uti", "cfu_threshold_source", "beoord_cat"
    }:
        return "Culture/symptoms"
    if col in {
        "analysis_include_primary", "analysis_exclusion_reason", "duplicate_role",
        "duplicate_of_participant_id", "duplicate_of_tp_lab", "allow_secondary_duplicate_qc",
        "duplicate_use_note", "genomics_expected_include", "genomics_exclusion_reason",
        "manual_curation_applied", "manual_curation_note", "manual_curation_source",
        "uricult_bridge_applied"
    }:
        return "QC/curation"
    return "Other"


def definition_for_column(col: str, group: str, gene_category: str | None) -> str:
    definitions = {
        "Participant_id": "Study participant identifier. This tells you which resident/person the row belongs to.",
        "tp_lab": "Timepoint label. Routine visits use labels such as T0, T1, and T2. UTI-related event samples use labels such as UTI-1.",
        "Episode_ID": "Unique episode identifier combining participant, timepoint/event type, and date. This is the safest single field to use for merging.",
        "Event_type": "Whether the row is from a scheduled routine sample or a UTI-event sample.",
        "Collection_Date": "Date the urine sample or episode was collected, stored as day/month/year text.",
        "Batch": "Sequencing or recruitment batch number.",
        "ST": "Sequence type: a bacterial lineage label for the E. coli isolate. Blank means no usable ST call.",
        "ST_source": "Where the sequence type came from, and whether the source passed the provider quality rule.",
        "ST_provider": "ST from the provider/RIVM SeqSphere source when available.",
        "ST_local": "ST from the local mlst run. A dash means no local ST was available.",
        "provider_PercGoodTargets": "Provider/RIVM SeqSphere quality measure: percentage of targets passing the provider quality checks. Provider calls were used when this was at least 95.",
        "provider_file": "Provider MLST file(s) from which the ST call was obtained.",
        "provider_batch_match": "Whether the provider MLST record matched the expected batch information.",
        "provider_assembler": "Which assembler's provider MLST result contributed to the ST call.",
        "UTI_Status": "Primary clinical classification used in the current analysis.",
        "Primary_Status": "Same primary UTI/Not_UTI classification repeated for convenience.",
        "UTI_binary": "Numeric version of UTI_Status: 1 = UTI and 0 = Not_UTI.",
        "Not_UTI_subgroup": "Subgroup label for rows that are Not_UTI.",
        "Infection_Status": "Current primary status label used by the VF-ready file.",
        "Infection_Status_legacy": "Older ASB/UTI/Negative label retained for comparison only.",
        "Infection_Status_old": "Same older ASB/UTI/Negative framing retained for audit/comparison.",
        "UTI_definition_version": "Name of the rule used to classify UTI versus Not_UTI.",
        "UTI_classification_confidence": "Confidence level for the clinical classification.",
        "UTI_classification_reason": "Plain reason the row was classified as UTI or Not_UTI.",
        "Status_Confidence_epi": "Episode-level confidence from the clinical source data.",
        "Sx_source_epi": "Where the symptom evidence came from.",
        "UTI_Label": "Original UTI label or identifier from source data when available.",
        "Urine_collection_method": "Human-readable urine collection method.",
        "urine_collection_method_raw": "Original urine collection method text from the source data.",
        "urine_collection_method_norm": "Standardized urine collection method used by the classification logic.",
        "catheter_rule": "Which symptom rule was used: non-catheter or indwelling-catheter logic.",
        "symptom_compatible_uti": "Whether recorded symptoms were compatible with the primary UTI rule.",
        "symptom_rule_met": "Specific symptom rule that was met, or that no symptom rule was met.",
        "local_urinary_symptom_any": "Whether any local urinary symptom was recorded.",
        "systemic_symptom_any": "Whether any systemic symptom was recorded.",
        "cfu_raw": "Original culture count text.",
        "cfu_raw_parsed": "Parsed culture count text used by the pipeline.",
        "cfu_ge_1e3": "Whether culture count was at least 1,000 CFU/mL.",
        "cfu_ge_1e4": "Whether culture count was at least 10,000 CFU/mL.",
        "cfu_ge_1e5": "Whether culture count was at least 100,000 CFU/mL.",
        "culture_supports_uti": "Whether culture results met the primary culture-support rule. In this VF-ready file, all rows meet culture support; UTI versus Not_UTI is separated by symptom compatibility.",
        "cfu_threshold_used_for_uti": "Culture threshold used by the primary UTI rule; the current file uses 1,000 CFU/mL.",
        "cfu_threshold_source": "How the culture threshold decision was derived.",
        "beoord_cat": "Original laboratory assessment category ('beoordeling' in Dutch), when available.",
        "analysis_include_primary": "Whether the episode is included in the primary analysis denominator.",
        "analysis_exclusion_reason": "Reason a row would be excluded from primary analysis.",
        "duplicate_role": "Whether this row is a duplicate or not.",
        "duplicate_of_participant_id": "If duplicate, the participant ID of the row it duplicates.",
        "duplicate_of_tp_lab": "If duplicate, the timepoint label of the row it duplicates.",
        "allow_secondary_duplicate_qc": "Whether a secondary duplicate was allowed for QC review.",
        "duplicate_use_note": "Note explaining duplicate handling.",
        "genomics_expected_include": "Whether the genomics row is expected and eligible for genomic analysis.",
        "genomics_exclusion_reason": "Reason the genomics row would be excluded.",
        "manual_curation_applied": "Whether a manual curation rule was applied.",
        "manual_curation_note": "Text note describing manual curation, if any.",
        "manual_curation_source": "Source of the manual curation decision.",
        "uricult_bridge_applied": "Whether a Uricult clinical event was bridged to a WGS row.",
        "vf_count_total": "Total number detected across all VF gene columns in this row.",
        "total_vf_count_all": "Same as vf_count_total: all VF genes detected.",
        "total_vf_count_curated": "Number detected among gene columns present in gene_map.csv when this file was built. This is not the later module/score-framework curated set.",
        "total_vf_count_upec_candidate": "Descriptive heuristic count of detected genes whose gene_map category/subcategory matched UPEC-relevant terms; this is not a validated score.",
        "total_vf_count_unassigned": "Number detected among gene columns absent from gene_map.csv; it uses the same gene set as cat_Unassigned_matrix.",
        "low_confidence_count": "Legacy placeholder fixed at zero in this file. Do not use it as the later score-framework low-confidence measure.",
        "cat_Adhesion_Fimbriae": "Number of detected genes in the adhesion/fimbriae category.",
        "cat_Unassigned": "Number detected among genes present in gene_map.csv but explicitly assigned to the Unassigned category.",
        "cat_Iron_acquisition": "Number of detected genes in the iron acquisition category.",
        "cat_Capsule_Surface": "Number of detected genes in the capsule/surface category.",
        "cat_Toxins": "Number of detected genes in the toxin category.",
        "cat_Invasion_Evasion": "Number of detected genes in the invasion/evasion category.",
        "cat_Unassigned_matrix": "Number detected among VF gene columns absent from gene_map.csv; it uses the same gene set as total_vf_count_unassigned.",
        "n_timepoints": "Number of distinct timepoints available for this participant in the VF-ready data.",
    }
    symptom_defs = {
        "flankpain_present": "Whether flank pain was recorded.",
        "dysuria_present": "Whether dysuria/painful urination was recorded.",
        "urgency_present": "Whether urinary urgency was recorded.",
        "frequency_present": "Whether urinary frequency was recorded.",
        "incontinence_present": "Whether incontinence was recorded.",
        "pus_present": "Whether pus/purulence was recorded.",
        "fever_present": "Whether fever was recorded.",
        "rigors_present": "Whether rigors/shaking chills were recorded.",
        "delirium_present": "Whether delirium/confusion was recorded.",
        "suprapubic_pain_present": "Whether suprapubic pain was recorded.",
        "other_sxs_present": "Whether other symptoms were recorded.",
    }
    definitions.update(symptom_defs)
    if group == "VF gene":
        category = gene_category or "not assigned in gene_map.csv"
        return f"Individual VFDB virulence-factor gene column. Category: {category}."
    return definitions.get(col, "Column retained from the source or derived during the VF-ready merge.")


def source_for_column(group: str, col: str) -> str:
    if group == "VF gene":
        return "Created from ABRicate/VFDB screening at >=80% identity and >=80% coverage; collapsed to 0/1 per episode."
    if group == "ST/MLST":
        return "Added from provider-preferred MLST output."
    if group in {"Clinical status", "Culture/symptoms"}:
        return "Added from the clinical status map and classification logic."
    if group == "VF summary":
        if col in {"vf_count_total", "total_vf_count_all"}:
            return "Sum of all VF gene 0/1 columns."
        if col == "total_vf_count_curated":
            return "Sum of VF columns present in gene_map.csv (script 22)."
        if col == "total_vf_count_upec_candidate":
            return "Heuristic gene_map category/subcategory term match (script 22)."
        if col in {"total_vf_count_unassigned", "cat_Unassigned_matrix"}:
            return "Sum of VF columns absent from gene_map.csv."
        if col == "low_confidence_count":
            return "Zero-filled legacy placeholder from script 22."
        if col == "n_timepoints":
            return "Distinct tp_lab count per Participant_id."
        return "Sum of VF columns in this gene_map.csv category."
    if group == "Identifier":
        return "Carried from metadata/status map and used to identify or merge rows."
    if group == "QC/curation":
        return "Carried from manual curation and analysis inclusion checks."
    return "Carried through from the merge process."


def codebook_rows(df: pd.DataFrame, gene_cols: set[str], gene_category: dict[str, str]) -> list[dict[str, str]]:
    rows = []
    seen = set()
    for col in df.columns:
        group = group_for_column(col, gene_cols)
        cat = gene_category.get(col)
        values = summarize_values(df[col], is_gene=group == "VF gene")
        missing = missing_text(df[col], col, group)
        rows.append({
            "Header": col,
            "Group": group,
            "Plain English meaning": definition_for_column(col, group, cat),
            "How to read filled-in values": values + " " + missing,
            "Source / note": source_for_column(group, col),
        })
        if col in seen:
            raise ValueError(f"Duplicate column in codebook generation: {col}")
        seen.add(col)
    if len(seen) != len(df.columns):
        raise ValueError("Header coverage mismatch")
    return rows


def mapped_gene_value(mapping: dict[str, str], gene: str, fallback: str) -> str:
    value = mapping.get(gene)
    if value is None or pd.isna(value) or str(value).strip() == "":
        return fallback
    return str(value)


def vf_gene_reference_rows(
    df: pd.DataFrame,
    gene_cols: set[str],
    gene_category: dict[str, str],
    gene_subcategory: dict[str, str],
) -> tuple[list[list[str]], list[list[str]], list[str]]:
    ordered_genes = [c for c in df.columns if c in gene_cols]
    preferred_category_order = [
        "Adhesion/Fimbriae",
        "Iron acquisition",
        "Toxins",
        "Capsule/Surface",
        "Invasion/Evasion",
        "Unassigned",
        "Not in gene_map.csv",
    ]

    by_category: dict[str, list[str]] = {category: [] for category in preferred_category_order}
    gene_rows = []
    for gene in ordered_genes:
        category = mapped_gene_value(gene_category, gene, "Not in gene_map.csv")
        subcategory = mapped_gene_value(gene_subcategory, gene, "Not mapped")
        by_category.setdefault(category, []).append(gene)
        gene_rows.append([
            gene,
            category,
            subcategory,
            "0 = gene not detected; 1 = gene detected.",
        ])

    def category_sort_key(category: str) -> tuple[int, str]:
        try:
            return (preferred_category_order.index(category), category)
        except ValueError:
            return (len(preferred_category_order), category)

    gene_rows.sort(key=lambda row: (category_sort_key(row[1]), row[0]))
    category_rows = [
        [category, str(len(genes)), ", ".join(sorted(genes))]
        for category, genes in sorted(by_category.items(), key=lambda item: category_sort_key(item[0]))
        if genes
    ]
    return category_rows, gene_rows, ordered_genes


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        st = doc.styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.15

    for style_name in ["List Bullet", "List Number"]:
        st = doc.styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.188)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def add_footer(section, text: str) -> None:
    p = section.footer.paragraphs[0]
    p.text = text
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def main() -> None:
    df = pd.read_csv(VF_READY, dtype=object)
    vf_pa = pd.read_csv(VF_PA, nrows=1, dtype=object)
    gene_cols = set([c for c in vf_pa.columns if c not in {"Participant_id", "tp_lab", "Episode_ID"}])
    gene_map_df = pd.read_csv(GENE_MAP, dtype=object)
    gene_category = dict(zip(gene_map_df["Gene"], gene_map_df["Category"]))
    gene_subcategory = dict(zip(gene_map_df["Gene"], gene_map_df["Subcategory"]))
    rows = codebook_rows(df, gene_cols, gene_category)
    vf_category_rows, vf_gene_rows, ordered_gene_cols = vf_gene_reference_rows(
        df, gene_cols, gene_category, gene_subcategory
    )

    gene_count = sum(1 for r in rows if r["Group"] == "VF gene")
    if len(rows) != len(df.columns):
        raise RuntimeError("Not every CSV header was included in the appendix.")
    if gene_count != 227:
        raise RuntimeError(f"Expected 227 VF gene columns, found {gene_count}.")
    if not all(c in df.columns for c in ["ST", "ST_source", "provider_PercGoodTargets", "UTI_Status", "Primary_Status"]):
        raise RuntimeError("Required ST or clinical status columns are missing.")

    gene_matrix = df[ordered_gene_cols].apply(pd.to_numeric, errors="raise")
    gene_values = set(pd.unique(gene_matrix.to_numpy().ravel()))
    if not gene_values.issubset({0, 1}) or gene_matrix.isna().any().any():
        raise RuntimeError(f"VF gene columns are not complete binary 0/1 data: {sorted(gene_values)}")
    if df["Episode_ID"].isna().any() or df["Episode_ID"].duplicated().any():
        raise RuntimeError("Episode_ID must be complete and unique before building the codebook.")
    vf_total = gene_matrix.sum(axis=1)
    if (vf_total != pd.to_numeric(df["vf_count_total"], errors="raise")).any():
        raise RuntimeError("vf_count_total does not equal the sum of the VF gene columns.")
    if not df["Primary_Status"].equals(df["UTI_Status"]):
        raise RuntimeError("Primary_Status and UTI_Status are not identical.")
    expected_binary = df["UTI_Status"].eq("UTI").astype(int)
    if (expected_binary != pd.to_numeric(df["UTI_binary"], errors="raise")).any():
        raise RuntimeError("UTI_binary is not aligned with UTI_Status.")
    provider_rows = df["ST_source"].eq("provider_qc95")
    provider_quality = pd.to_numeric(df.loc[provider_rows, "provider_PercGoodTargets"], errors="coerce")
    if provider_quality.isna().any() or (provider_quality < 95).any():
        raise RuntimeError("A provider_qc95 row lacks provider_PercGoodTargets >= 95.")

    n_uti = int(df["UTI_Status"].eq("UTI").sum())
    n_not_uti = int(df["UTI_Status"].eq("Not_UTI").sum())
    st_source_counts = df["ST_source"].value_counts(dropna=False)
    n_provider = int(st_source_counts.get("provider_qc95", 0))
    n_local_fallback = int(
        sum(count for source, count in st_source_counts.items() if str(source).startswith("local_fallback"))
    )
    n_missing_st = int(df["ST"].isna().sum())
    culture_supported = int(
        df["culture_supports_uti"].astype(str).str.strip().str.lower().eq("true").sum()
    )
    verification_date = date.today().strftime("%d %B %Y").lstrip("0")

    doc = Document()
    configure_styles(doc)
    add_footer(doc.sections[0], "vf_analysis_ready.csv codebook")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Codebook for vf_analysis_ready.csv: Virulence Factor and Sequence Type Dataset")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Lay-friendly guide to the merged VF, ST, and clinical analysis file")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    add_body(
        doc,
        "Audience: people who need to understand and use the dataset without reading the R scripts."
    )
    add_body(
        doc,
        f"Dataset summarized: {VF_READY.name}. Current file has {len(df):,} rows, "
        f"{len(df.columns)} columns, and {gene_count} individual virulence-factor gene columns."
    )
    add_body(doc, f"Codebook verified against the current primary analysis files: {verification_date}.")
    add_body(
        doc,
        "Best merge key: Episode_ID. If that is not available in another dataset, use Participant_id plus tp_lab and check carefully for duplicates."
    )
    doc.add_page_break()

    add_heading(doc, "Overview", 1)
    add_body(
        doc,
        "This file is an episode-level table. Each row represents one E. coli isolate/episode linked to a participant and a timepoint. "
        "The table combines genome-screening results, sequence type, clinical classification, culture/symptom evidence, and summary counts."
    )
    add_note_box(
        doc,
        "Plain-language reading rule",
        "Start with the identifier columns to see who and when the row is from. Then read the VF gene columns as detected/not detected, the ST columns as bacterial lineage information, and the status columns as the clinical UTI/Not_UTI interpretation."
    )
    add_heading(doc, "Quick Start: What To Use For Analysis", 2)
    quick_start_rows = [
        ["Main datafile", "vf_analysis_ready.csv", "Use this as the main analysis file because it already combines episode identifiers, ST calls, VF genes, clinical status, culture/symptom fields, and VF summary counts."],
        ["Best merge key", "Episode_ID", "Safest single field for merging with another episode-level dataset."],
        ["Backup merge key", "Participant_id + tp_lab", "Use only if Episode_ID is not available in the other dataset; check duplicates and dates carefully."],
        ["Sequence type", "ST", "Final sequence type used in the analysis. ST means bacterial lineage label."],
        ["ST provenance", "ST_source; provider_PercGoodTargets", f"Provider/RIVM SeqSphere calls with at least 95 percent good targets were preferred. Current sources: {n_provider} provider, {n_local_fallback} local fallback, and {n_missing_st} missing."],
        ["Clinical grouping", "UTI_Status or Primary_Status", f"Use the current UTI versus Not_UTI classification ({n_uti} UTI and {n_not_uti} Not_UTI). Legacy ASB/UTI/Negative fields are for comparison only."],
        ["Individual VF genes", "Gene-name columns in Appendix B", "Each uses 1 = detected and 0 = not detected after ABRicate/VFDB screening at >=80% identity and >=80% coverage."],
        ["VF summaries", "vf_count_total and cat_* columns", "Use for descriptive counts overall and by the simple gene_map.csv categories. Later module/score outputs use a separate annotation framework."],
        ["Heatmap", "variable_gene_heatmap.png", "Visual summary only. Do not extract or merge data from the image; use vf_analysis_ready.csv."],
        ["Main cautions", "Interpret with caveats", "Detected does not mean active or expressed; missing ST means no usable ST call; VF counts are descriptive and not a validated severity score."],
    ]
    add_table(doc, ["Need", "Use this column/file", "How to explain it"], quick_start_rows, [1.35, 2.15, 4.0], font_size=8)

    add_note_box(
        doc,
        "Current dataset-specific status context",
        f"All {culture_supported} rows in this sequenced VF-ready subset meet the culture-support rule. The primary split is therefore symptom-based within this subset: {n_uti} rows meet the compatible symptom rule and are UTI; {n_not_uti} do not meet it and are Not_UTI (bacteriuria_not_UTI). Not_UTI here does not mean culture-negative."
    )

    add_heading(doc, "How this file was produced", 2)
    for item in [
        "One canonical E. coli assembly was selected for each included participant-timepoint/episode after genomics QC.",
        "ABRicate screened each selected assembly against VFDB using at least 80% sequence identity and at least 80% gene coverage.",
        "Gene hits were collapsed to one binary value per episode: 1 = detected and 0 = not detected, producing vf_pa_all.csv.",
        "Script 22 joined the binary VF matrix to the primary clinical status map and the provider-preferred MLST table.",
        "Provider/RIVM SeqSphere ST calls with provider_PercGoodTargets >= 95 were preferred; local MLST was used only as an explicitly labelled fallback.",
        "The final step added VF totals, simple gene_map.csv category counts, provenance fields, and QC/curation fields to create vf_analysis_ready.csv.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "The three most important ideas", 2)
    for item in [
        "A VF gene value of 1 means the gene was detected at >=80% identity and >=80% coverage. A value of 0 means it was not detected at those thresholds.",
        "ST means sequence type, a bacterial lineage label. It helps compare isolates, but it does not by itself prove that two isolates are the same strain.",
        "Primary_Status and UTI_Status are the current clinical labels. The older ASB/UTI/Negative labels are retained only for comparison.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Snapshot of the current file", 2)
    diag_text = DIAGNOSTICS.read_text()
    diag_rows = [
        ["Codebook verification date", verification_date],
        ["Diagnostics timestamp", diagnostic_value(diag_text, "Timestamp")],
        ["Rows", f"{len(df):,}"],
        ["Participants", f"{df['Participant_id'].nunique():,}"],
        ["VF gene columns", str(gene_count)],
        ["Primary UTI rows", str(n_uti)],
        ["Primary Not_UTI rows", str(n_not_uti)],
        ["Rows with ST", str(df["ST"].notna().sum())],
        ["Rows without ST", str(n_missing_st)],
        ["Distinct STs", str(df["ST"].nunique(dropna=True))],
        ["ST source breakdown", f"provider_qc95={n_provider}; local fallback={n_local_fallback}; missing={n_missing_st}"],
        ["Status rows not in VF", diagnostic_value(diag_text, "Status rows NOT in VF")],
        ["Genes in matrix not in gene_map", diagnostic_value(diag_text, "Genes in matrix but NOT in gene_map")],
    ]
    add_table(doc, ["Item", "Current value"], diag_rows, [2.2, 4.3], font_size=9)

    add_heading(doc, "How to read one row", 1)
    example = df[df["UTI_Status"] == "UTI"].iloc[0]
    example_rows = [
        ["Participant and episode", f"Participant {clean_value(example['Participant_id'])}, {clean_value(example['tp_lab'])}, Episode_ID {clean_value(example['Episode_ID'])}"],
        ["Clinical status", f"{clean_value(example['UTI_Status'])}; reason: {clean_value(example['UTI_classification_reason'])}"],
        ["Sequence type", f"ST {clean_value(example['ST'])}; source: {clean_value(example['ST_source'])}; provider good targets: {clean_value(example['provider_PercGoodTargets'])}"],
        ["VF burden", f"{clean_value(example['vf_count_total'])} total detected VF genes"],
        ["Category examples", f"Adhesion/fimbriae: {clean_value(example['cat_Adhesion_Fimbriae'])}; toxins: {clean_value(example['cat_Toxins'])}; iron acquisition: {clean_value(example['cat_Iron_acquisition'])}"],
    ]
    add_table(doc, ["Question", "How this example row reads"], example_rows, [2.0, 4.5], font_size=9)
    add_body(
        doc,
        "In words: this row is a UTI-event E. coli isolate from one participant/timepoint. It has an ST call, a measured VF gene profile, and clinical evidence supporting the primary UTI classification."
    )

    add_heading(doc, "Column groups", 1)
    group_rows = [
        ["Identifier", "Columns that tell you who, when, and which episode the row belongs to.", "Use Episode_ID for merging where possible."],
        ["VF gene", "Individual gene-name columns from the ABRicate/VFDB screen.", "0 = not detected; 1 = detected at >=80% identity and >=80% coverage."],
        ["ST/MLST", "Sequence type and where the sequence type call came from.", "ST is lineage context, not a same-strain proof."],
        ["Clinical status", "Primary UTI/Not_UTI classification and legacy status labels.", "Use UTI_Status or Primary_Status for the current analysis."],
        ["Culture/symptoms", "Culture count, symptom, catheter, and clinical rule fields.", "These explain why a row was called UTI or Not_UTI."],
        ["VF summary", "Counts of detected genes overall and by biological category.", "Useful for summaries; not a validated severity score."],
        ["QC/curation", "Inclusion, duplicate, and manual-curation fields.", "Mostly audit fields; blanks often mean not applicable."],
    ]
    add_table(doc, ["Column group", "What it contains", "How to use it"], group_rows, [1.45, 3.25, 2.8], font_size=8)

    add_heading(doc, "Value legend", 1)
    legend_rows = [
        ["0 in a VF gene column", "Gene was not detected at the >=80% identity and >=80% coverage screening thresholds."],
        ["1 in a VF gene column", "Gene was detected by ABRicate/VFDB at >=80% identity and >=80% coverage."],
        ["UTI", "Primary clinical rule classified this episode as UTI."],
        ["Not_UTI", "Primary clinical rule did not classify this episode as UTI. In this VF-ready subset, this means culture-supported bacteriuria without a compatible symptom rule, not a negative culture."],
        ["1 in UTI_binary", "The row is classified as UTI."],
        ["0 in UTI_binary", "The row is classified as Not_UTI."],
        ["ASB / Negative / UTI in legacy columns", "Older clinical framing retained for comparison; do not use as the main current status unless specifically doing legacy comparison."],
        ["provider_qc95", "Provider/RIVM ST call used, with at least 95 percent good targets."],
        ["local_fallback_provider_missing", "Local MLST was used because provider ST was missing."],
        ["missing ST", "No usable sequence type call was available."],
        ["TRUE", "Yes, condition met, recorded, or included."],
        ["FALSE", "No, condition not met, not recorded, or not included."],
        ["blank/NA", "Usually not recorded, unavailable, or not applicable; exact meaning depends on the column."],
    ]
    add_table(doc, ["Value", "Plain-English meaning"], legend_rows, [2.0, 4.5], font_size=9)

    add_heading(doc, "Important caveats", 1)
    for item in [
        "Detected does not mean the gene is switched on or causing disease; it only means the sequence was found at the >=80% identity and >=80% coverage thresholds.",
        "Not detected does not prove absolute absence; it means the gene was not found at those thresholds.",
        "ST helps describe bacterial lineage. Same ST can support similarity, but it does not prove same strain by itself.",
        "Category counts such as cat_Toxins and cat_Adhesion_Fimbriae are descriptive counts, not proven virulence severity scores.",
        "The cat_* columns and total_vf_count_curated use the simple gene_map.csv grouping used to build vf_analysis_ready.csv. Later module/score outputs use gene_module_map.csv and a separate curation framework, so similarly named counts should not be treated as interchangeable.",
        "cat_Unassigned counts genes that are present in gene_map.csv but labelled Unassigned. cat_Unassigned_matrix and total_vf_count_unassigned count genes absent from gene_map.csv.",
        "low_confidence_count is a legacy zero-filled placeholder in this file; later score-framework low-confidence annotations are separate.",
        "Repeated rows from the same participant are related observations, so they should not be treated as fully independent in statistical analyses.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Major column groups in plain language", 1)
    plain_sections = [
        ("Identifiers", "Use these first. They tell you which person, timepoint, episode, date, and batch the row describes."),
        ("VF gene columns", "These are the many gene-name columns near the start of the file. Each asks whether the VFDB gene was detected by ABRicate at >=80% identity and >=80% coverage."),
        ("ST/MLST columns", "These describe the sequence type and where the ST came from. ST is useful lineage context, especially when comparing isolates across time."),
        ("Clinical/status columns", "These show the current UTI/Not_UTI classification and retain older labels for audit. For current analysis, use UTI_Status or Primary_Status."),
        ("Culture and symptoms", "These fields explain the clinical logic: urine collection method, catheter rule, symptoms, culture count thresholds, and whether culture supported UTI."),
        ("VF summaries", "These columns count detected genes overall and within the simple gene_map.csv groupings such as adhesion/fimbriae, toxins, and iron acquisition. They are descriptive and are separate from later module/score-framework endpoints."),
        ("QC and curation", "These fields document whether rows were included, excluded, duplicated, manually curated, or bridged between clinical and WGS labels."),
    ]
    for heading, text in plain_sections:
        add_heading(doc, heading, 2)
        add_body(doc, text)

    doc.add_page_break()
    appendix = doc.add_section(WD_SECTION.NEW_PAGE)
    appendix.footer.is_linked_to_previous = False
    appendix.orientation = WD_ORIENT.LANDSCAPE
    appendix.page_width = Inches(11)
    appendix.page_height = Inches(8.5)
    appendix.top_margin = Inches(0.55)
    appendix.bottom_margin = Inches(0.55)
    appendix.left_margin = Inches(0.55)
    appendix.right_margin = Inches(0.55)
    add_footer(appendix, "Appendix: full header reference")

    add_heading(doc, "Appendix A: non-gene header reference", 1)
    add_body(
        doc,
        "This appendix lists the non-gene headers in vf_analysis_ready.csv. Individual VF gene columns are separated into Appendix B so the lookup tables stay readable."
    )

    order = [
        "Identifier", "Clinical status", "Culture/symptoms",
        "QC/curation", "ST/MLST", "VF summary", "Other"
    ]
    rows_by_group = {g: [r for r in rows if r["Group"] == g] for g in order}
    widths = [1.55, 1.1, 3.1, 3.25, 1.1]
    headers = ["Header", "Group", "Plain English meaning", "How to read filled-in values", "Source / note"]
    appended_headers = []
    for group in order:
        group_rows = rows_by_group.get(group, [])
        if not group_rows:
            continue
        add_heading(doc, group, 2)
        table_rows = [[r[h] for h in headers] for r in group_rows]
        add_table(doc, headers, table_rows, widths, font_size=7)
        appended_headers.extend([r["Header"] for r in group_rows])
        add_body(doc, "")

    add_heading(doc, "Appendix B: VF gene category reference", 1)
    add_body(
        doc,
        "This table lists every individual virulence-factor gene column exactly once, grouped using gene_map.csv, the simple map used to build the cat_* columns in this file. Each gene uses the same coding: 1 means detected by ABRicate/VFDB at >=80% identity and >=80% coverage; 0 means not detected at those thresholds. These categories are separate from the later gene_module_map.csv score/module framework. Use Word's Find command to locate a specific gene name."
    )
    add_table(
        doc,
        ["Category", "Number of VF gene columns", "VF gene columns in this category"],
        vf_category_rows,
        [1.8, 1.4, 6.3],
        font_size=7,
    )
    appended_headers.extend([r[0] for r in vf_gene_rows])

    if set(appended_headers) != set(df.columns):
        missing = set(df.columns) - set(appended_headers)
        extra = set(appended_headers) - set(df.columns)
        raise RuntimeError(f"Appendix mismatch. Missing={missing}; extra={extra}")
    if len(appended_headers) != len(set(appended_headers)):
        raise RuntimeError("Appendix contains duplicate headers.")
    if set(ordered_gene_cols) != {r[0] for r in vf_gene_rows}:
        raise RuntimeError("VF gene reference table does not match the VF gene columns.")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)
    print(f"headers={len(appended_headers)} vf_gene_columns={gene_count}")


if __name__ == "__main__":
    main()
