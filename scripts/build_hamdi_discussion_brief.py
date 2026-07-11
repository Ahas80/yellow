#!/usr/bin/env python3
"""Build a meeting-ready discussion brief for the Hamdi thesis audit.

Outputs:
  - results/thesis_audit/hamdi_thesis_discussion_brief_today.docx
  - results/thesis_audit/hamdi_thesis_discussion_brief_today.md

This document consolidates the audit plan, independent audit, hypothesis
summary, and calculation/filter reconciliation into a short internal brief.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "results" / "thesis_audit"
DOCX_OUT = OUT_DIR / "hamdi_thesis_discussion_brief_today.docx"
MD_OUT = OUT_DIR / "hamdi_thesis_discussion_brief_today.md"


# standard_business_brief tokens
FONT = "Calibri"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x0B, 0x25, 0x45)
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_FILL = "F2F4F7"
CALL_FILL = "F4F6F9"
CAUTION_FILL = "FFF7D6"
RISK_FILL = "FDECEC"
OK_FILL = "EAF5EF"
BORDER = "C9D1D9"


def set_run_font(run, size: float | None = None, color: RGBColor | None = None,
                 bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_paragraph(paragraph, before=0, after=6, line=1.10,
                    align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_para(doc, text="", style: str | None = None, size=11, color=None,
             bold=False, italic=False, before=0, after=6, line=1.10,
             align: WD_ALIGN_PARAGRAPH | None = None):
    p = doc.add_paragraph(style=style)
    style_paragraph(p, before=before, after=after, line=line, align=align)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_rich_para(doc, parts, before=0, after=6, line=1.10, style=None):
    p = doc.add_paragraph(style=style)
    style_paragraph(p, before=before, after=after, line=line)
    for text, attrs in parts:
        run = p.add_run(text)
        set_run_font(
            run,
            size=attrs.get("size", 11),
            color=attrs.get("color"),
            bold=attrs.get("bold"),
            italic=attrs.get("italic"),
        )
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    if level == 1:
        style_paragraph(p, before=16, after=8, line=1.10)
        size, color = 16, BLUE
    elif level == 2:
        style_paragraph(p, before=12, after=6, line=1.10)
        size, color = 13, BLUE
    else:
        style_paragraph(p, before=8, after=4, line=1.10)
        size, color = 12, DARK_BLUE
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=True)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = add_para(doc, text, style=style, after=4, line=1.167)
    p.paragraph_format.left_indent = Inches(0.5 if level == 0 else 0.75)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    return p


def add_numbered(doc, text):
    p = add_para(doc, text, style="List Number", after=4, line=1.167)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    return p


def add_callout(doc, label, text, fill=CALL_FILL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, BORDER)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_paragraph(p, before=0, after=0, line=1.10)
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=10.5, color=INK, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    add_para(doc, "", after=4)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table, BORDER)
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_shading(hdr[idx], LIGHT_FILL)
        hdr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr[idx].paragraphs[0]
        style_paragraph(p, before=0, after=0, line=1.10)
        r = p.add_run(h)
        set_run_font(r, size=9.5, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            p = cells[idx].paragraphs[0]
            style_paragraph(p, before=0, after=0, line=1.08)
            r = p.add_run(str(val))
            set_run_font(r, size=9.2, color=RGBColor(0, 0, 0))
    add_para(doc, "", after=4)
    return table


def build_markdown() -> str:
    return """# Independent Scientific Audit Discussion Brief

**Subject:** Hamdi Hersi thesis results vs current rUTI/YELLOW RoUTIne analyses  
**Prepared for:** Discussion today  
**Date:** 7 July 2026  
**Purpose:** Internal scientific-validity and reproducibility discussion. This is not an allegation about intent.

## 60-second opening

I have compared the thesis claims against my current generated analyses. The descriptive genomic backbone is partly corroborated: the isolate count, batch count, ST diversity and broad ST131/ST73 trends line up. The concern is not that the whole dataset is wrong. The concern is that several headline conclusions appear to depend on different denominators, event-label definitions, wgMLST/SNP non-equivalence, grey-zone pair handling, and under-specified antibiotic/repeated-measures modelling. I think we should reconcile the row-level wgMLST pair table, transition definitions, and model-ready antibiotic data before treating the replacement, ASB-to-UTI and causal antibiotic claims as verified.

## Bottom line

The fair position is: **partly corroborated descriptively, not yet reproducible for the main transition and causal claims.**

Avoid a binary “legit/not legit” verdict. Classify claims as verified, independently corroborated, discrepant, methodologically unsupported, or not checkable with available data.

## What matches or broadly makes sense

- 556 WGS/VF isolate rows are present in the current analysis.
- Six sequencing batches are represented.
- There are 83 unique typed STs.
- Median WGS-ready timepoints per resident is 3, IQR 2-4, range 1-9.
- ST131 decreases over time and ST73 increases over time.
- These agreements support dataset overlap and broad lineage trends, but they do not validate wgMLST, recurrence, antibiotic, or causal claims.

## Core filter reconciliation

| Topic | Thesis appears to use | Current pipeline uses | Why it matters |
|---|---|---|---|
| Resident universe | 167 genomic residents; 166 linked to Castor; one UNKNOWN | 167 full status-map participants, 166 primary clinical participants, 162 VF/WGS-ready linked participants | Same isolate count can coexist with different participant denominators. |
| UTI/ASB labels | Routine timepoints treated as ASB; UTI-* treated as UTI episodes | `UTI_Status` based on culture support plus compatible symptoms; `Event_type` kept separate | Her Routine→UTI-event analysis is not identical to your Not_UTI→UTI analysis. |
| Pair universe | 368 consecutive wgMLST pairs, possibly after comparable-loci exclusions | 394 WGS/VF-linked consecutive transitions after `cohort == "all"` | Difference may be wgMLST filtering, but needs row-level attrition. |
| Strain identity | wgMLST ≤25 alleles same strain; 26-100 grey zone; >100 replacement | SNP ≤25 strict same strain; same-ST >25 SNPs is not automatically replacement | SNP and wgMLST are not interchangeable. |
| Antibiotic models | Castor antibiotic windows linked to pairs | No antibiotic data in current generated outputs | Do not reconstruct from proxies. |

## Current independent anchors

- Primary clinical episodes: 583 rows from 166 participants.
- Full clinical status map: 585 rows from 167 participants.
- Canonical WGS/VF rows: 556 rows from 162 participants.
- Repeated-measures VF subset: 538 episodes from 144 participants.
- Consecutive WGS/VF transitions: 394 from 144 participants after `cohort == "all"`.
- Strict SNP same strain: 116/394 (29.4%).
- Replacement likely: 106/394 (26.9%).
- Same lineage but not same strain by SNP: 172/394 (43.7%).
- Routine→UTI-event transitions: 24 from 23 participants.
- Routine→primary-UTI transitions: 9 from 9 participants.
- All primary Not_UTI→UTI transitions: 10 from 10 participants.

## What does not make sense yet

### 1. Grey-zone handling

The thesis reports 231 persistent, 26 grey-zone and 111 replacement pairs. That sums to 368. Later the Cox model reports 137 replacement events, which equals 111 + 26. This suggests grey-zone pairs may have been described as ambiguous but counted as replacement in modelling.

### 2. Replacement threshold shifts

Methods say replacement is >25 wgMLST alleles. Results/Figure split 26-100 as grey zone and >100 as replacement. Discussion says grey-zone pairs may still be classed as replacement. This needs one explicit rule.

### 3. ASB→UTI denominator mismatch

The thesis reports 34 ASB→UTI pairs across 26 residents, but another transition partition uses only 26 ASB→UTI pairs. Your event-label analysis finds 24 Routine→UTI-event transitions; your current primary definition finds only 9 Routine→primary-UTI transitions.

### 4. Post-UTI return not reproduced

The thesis reports 37 post-UTI pairs and 25/37 returns. Your current complete triple analysis finds 20 event-label triples and 8 primary-UTI triples. Exact replication needs her row-level pre-ASB/UTI/post-UTI table.

### 5. Antibiotic model denominator

The thesis says antibiotic exposure data are available for 337/368 pairs, but the Cox model is reported with n=368. Missing-exposure handling is not described.

### 6. Repeated observations

Pair-level tests are not independent because residents contribute repeated pairs. The thesis does not clearly report clustering, frailty, mixed models, GEE, robust standard errors, or a recurrent-event survival framework.

### 7. Assembly/wgMLST concern

The thesis says Flye was selected because it was more accurate than Longcycler, but the cited Landman et al. benchmark reported Longcycler/Miniasm as better for wgMLST allele calling. Flye’s E. coli 95th-percentile wgMLST discrepancy was 26.1 alleles, larger than the thesis’s 25-allele same-strain threshold. The cited scheme reports 4,503 loci, while the thesis reports 4,512.

### 8. Genetic drift wording

The claim of “approximately one allele per month” is not a formal rate estimate unless modelled against elapsed time with uncertainty. The subset is also selected for stable persistence, so saying no one crossed the threshold is partly built into the selection.

### 9. Recurrence wording

The thesis says residents were selected for re-emergence of the original strain, then describes remaining cases as cases where the original strain did not return. That wording needs correction and a case table.

### 10. Causal language

Phrases such as “primary driver,” “protective colonisation,” and “strengthened causal inference” are too strong for the methods as written. The safer wording is association or hypothesis-generating evidence.

## How to say it today

Use this framing:

> I am not trying to infer intent or make an accusation. I am trying to determine which claims are reproducible from the current data and which require the original wgMLST, antibiotic, and model-ready tables.

Then:

> The basic genomic dataset appears broadly aligned with my analysis. The main concerns are denominator definitions, UTI event labels versus primary clinical UTI status, treatment of grey-zone pairs, and under-specified antibiotic/repeated-measures modelling.

And:

> Before accepting the headline claims, I think we need the row-level wgMLST pair-distance table, the transition case table, and the model-ready antibiotic dataset.

## Say this / avoid this

| Say this | Avoid this |
|---|---|
| “The descriptive genomic results are partly corroborated.” | “The thesis is fake.” |
| “The headline transition and causal claims are not yet reproducible from available materials.” | “The results are illegitimate.” |
| “The differences may be due to denominator and clinical-definition choices.” | “She must have manipulated the data.” |
| “Grey-zone handling needs explicit reconciliation.” | “The model is definitely wrong.” |
| “Antibiotic claims should be treated as associations unless causal design is supplied.” | “Antibiotics definitely do not matter.” |

## Questions to ask

1. Can we see the row-level wgMLST isolate and pair-distance table, including comparable loci and missing loci?
2. Which exact E. coli wgMLST scheme version was used, and why does the thesis report 4,512 loci when the cited paper reports 4,503?
3. Were the 26 grey-zone pairs excluded, treated as uncertain, or counted as replacement in the trajectory and Cox analyses?
4. Why are there 368 consecutive wgMLST pairs but 394 WGS/VF-linked transitions in my current pipeline?
5. Why does the thesis report 34 ASB→UTI pairs but also use 26 ASB→UTI pairs in another transition partition?
6. How were UTI episodes defined: event label, Castor UTI, culture/symptom rule, or sequenced UTI isolate?
7. How were repeated pairs from the same resident handled statistically?
8. For the Cox model, what was the time origin, event definition, censoring rule, clustering method, proportional-hazards diagnostic and missing-exposure handling?
9. Can we see the Castor antibiotic extract with indication, drug, start/end dates and linkage rules?
10. Can causal language be softened unless a causal model/design is provided?

## Evidence still needed

- wgMLST pair-distance table and isolate-level allele profiles.
- Assembly/QC manifest with assembler choice, coverage, polishing and QC provenance.
- Exact SeqSphere scheme name/version/date and target export.
- Row-level ASB/UTI transition and recurrence tables.
- Castor antibiotic extract and model-ready antibiotic variables.
- Model formulas, code, package/session info and missing supplementary tables.

## Final position

The scientific position to take is:

> The thesis is not fully verifiable from the available materials. Basic dataset size, batch count, ST diversity and broad ST trends are supported. The exact persistence, recurrence, ASB-to-UTI, post-UTI return and antibiotic estimates remain unverified, and several methodological inconsistencies materially weaken causal conclusions. The next step is targeted clarification and source tables, not an inference about intent.
"""


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ["List Bullet", "List Bullet 2", "List Number"]:
        if style_name in styles:
            st = styles[style_name]
            st.font.name = FONT
            st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
            st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
            st.font.size = Pt(11)
            st.paragraph_format.space_after = Pt(4)
            st.paragraph_format.line_spacing = 1.167

    # Simple memo masthead title block.
    add_para(doc, "INTERNAL SCIENTIFIC DISCUSSION BRIEF", size=10.5, color=MUTED, bold=True, after=4)
    title = add_para(
        doc,
        "Hamdi Hersi Thesis Results vs Current rUTI Analyses",
        size=22,
        color=RGBColor(0, 0, 0),
        bold=True,
        after=3,
        line=1.05,
    )
    subtitle = add_para(
        doc,
        "Independent audit consolidation for discussion today",
        size=12.5,
        color=MUTED,
        after=10,
    )
    add_rich_para(
        doc,
        [
            ("Date: ", {"bold": True, "color": INK}),
            ("7 July 2026", {}),
            ("    Purpose: ", {"bold": True, "color": INK}),
            ("scientific validity and reproducibility; not an allegation about intent", {}),
        ],
        after=10,
    )

    add_callout(
        doc,
        "Meeting bottom line",
        "The descriptive genomic backbone is partly corroborated, but the main transition, recurrence, post-UTI return and antibiotic/causal claims are not yet reproducible from the available materials.",
        fill=CAUTION_FILL,
    )

    add_heading(doc, "60-second opening", 1)
    add_para(
        doc,
        "I have compared the thesis claims against my current generated analyses. The descriptive genomic backbone is partly corroborated: the isolate count, batch count, ST diversity and broad ST131/ST73 trends line up. The concern is not that the whole dataset is wrong. The concern is that several headline conclusions appear to depend on different denominators, event-label definitions, wgMLST/SNP non-equivalence, grey-zone pair handling, and under-specified antibiotic/repeated-measures modelling. I think we should reconcile the row-level wgMLST pair table, transition definitions, and model-ready antibiotic data before treating the replacement, ASB-to-UTI and causal antibiotic claims as verified.",
    )

    add_heading(doc, "What is supported", 1)
    for item in [
        "556 WGS/VF isolate rows are present in the current analysis.",
        "Six sequencing batches are represented.",
        "There are 83 unique typed sequence types.",
        "Median WGS-ready timepoints per resident is 3, IQR 2-4, range 1-9.",
        "ST131 decreases over time and ST73 increases over time.",
        "These agreements support dataset overlap and broad lineage trends; they do not validate wgMLST, recurrence, antibiotic, or causal claims.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    add_heading(doc, "Core filter reconciliation", 1)
    add_table(
        doc,
        ["Topic", "Thesis appears to use", "Current pipeline uses", "Why it matters"],
        [
            [
                "Resident universe",
                "167 genomic residents; 166 linked to Castor; one UNKNOWN",
                "167 full status-map participants, 166 primary clinical participants, 162 VF/WGS-ready linked participants",
                "Same isolate count can coexist with different participant denominators.",
            ],
            [
                "UTI/ASB labels",
                "Routine timepoints treated as ASB; UTI-* treated as UTI episodes",
                "UTI_Status based on culture support plus compatible symptoms; Event_type kept separate",
                "Routine->UTI-event is not the same as Not_UTI->UTI.",
            ],
            [
                "Pair universe",
                "368 consecutive wgMLST pairs, possibly after comparable-loci exclusions",
                "394 WGS/VF-linked consecutive transitions after cohort == all",
                "Difference may be wgMLST filtering, but needs row-level attrition.",
            ],
            [
                "Strain identity",
                "wgMLST <=25 alleles same strain; 26-100 grey zone; >100 replacement",
                "SNP <=25 strict same strain; same-ST >25 SNPs is not automatically replacement",
                "SNP and wgMLST are not interchangeable.",
            ],
        ],
        [1500, 2600, 2750, 2510],
    )

    add_heading(doc, "Current independent anchors", 1)
    anchors = [
        ("Primary clinical episodes", "583 rows from 166 participants"),
        ("Full clinical status map", "585 rows from 167 participants"),
        ("Canonical WGS/VF rows", "556 rows from 162 participants"),
        ("Repeated-measures VF subset", "538 episodes from 144 participants"),
        ("Consecutive WGS/VF transitions", "394 from 144 participants after cohort == all"),
        ("Strict SNP same strain", "116/394 (29.4%)"),
        ("Replacement likely", "106/394 (26.9%)"),
        ("Same lineage but not same strain by SNP", "172/394 (43.7%)"),
        ("Routine->UTI-event transitions", "24 from 23 participants"),
        ("Routine->primary-UTI transitions", "9 from 9 participants"),
        ("All primary Not_UTI->UTI transitions", "10 from 10 participants"),
    ]
    add_table(doc, ["Anchor", "Current result"], anchors, [3300, 6060])

    add_heading(doc, "What does not make sense yet", 1)
    concerns = [
        ("Grey-zone handling", "231 persistent + 26 grey-zone + 111 replacement sums to 368, but the later Cox model reports 137 replacement events, which equals 111 + 26. This suggests grey-zone pairs may have been presented as ambiguous but counted as replacement."),
        ("Replacement threshold shifts", "Methods imply >25 alleles is replacement. Results use 26-100 as grey zone and >100 as replacement. Discussion then suggests grey-zone pairs were still classed as replacement."),
        ("ASB-to-UTI denominators", "The thesis reports 34 ASB-to-UTI pairs across 26 residents, but another transition partition uses 26 ASB-to-UTI pairs. Current event-label analysis finds 24; current primary clinical definition finds 9 Routine->primary-UTI transitions."),
        ("Post-UTI return", "The thesis reports 37 post-UTI pairs and 25/37 returns. Current complete triple analysis finds 20 event-label triples and 8 primary-UTI triples. Exact replication needs her case table."),
        ("Antibiotic model denominator", "Antibiotic exposure is reported for 337/368 pairs, but the Cox model is reported with n=368. Missing-exposure handling is not described."),
        ("Repeated observations", "Residents contribute repeated pairs. The thesis does not clearly report clustering, frailty, mixed models, GEE, robust standard errors, or recurrent-event survival methods."),
        ("Assembly/wgMLST concern", "Flye was justified as more accurate, but Landman et al. reported Longcycler/Miniasm as better for wgMLST allele calling. Flye’s E. coli 95th-percentile discrepancy was 26.1 alleles, above the thesis’s 25-allele threshold. Scheme loci also differ: 4,503 in the cited benchmark vs 4,512 in the thesis."),
        ("Genetic drift wording", "Approximately one allele per month is not a formal rate estimate unless modelled against elapsed time with uncertainty. The stable-persistence subset creates circularity for claims about staying below the threshold."),
        ("Recurrence wording", "The thesis selects residents with original-strain re-emergence, then says remaining cases are where the original strain did not return. That needs a corrected case definition."),
        ("Causal language", "Terms such as primary driver, protective colonisation, and strengthened causal inference are too strong for the described observational analyses."),
    ]
    for label, text in concerns:
        add_rich_para(doc, [(f"{label}: ", {"bold": True, "color": DARK_BLUE}), (text, {})], after=5)

    add_heading(doc, "How to say it today", 1)
    add_callout(
        doc,
        "Use this framing",
        "I am not trying to infer intent or make an accusation. I am trying to determine which claims are reproducible from the current data and which require the original wgMLST, antibiotic, and model-ready tables.",
        fill=OK_FILL,
    )
    add_callout(
        doc,
        "Then say",
        "The basic genomic dataset appears broadly aligned with my analysis. The main concerns are denominator definitions, UTI event labels versus primary clinical UTI status, treatment of grey-zone pairs, and under-specified antibiotic/repeated-measures modelling.",
        fill=OK_FILL,
    )
    add_callout(
        doc,
        "Ask for",
        "The row-level wgMLST pair-distance table, transition case table, assembly/QC manifest, and model-ready antibiotic dataset before accepting the headline claims.",
        fill=OK_FILL,
    )

    doc.add_page_break()
    add_heading(doc, "Say this / avoid this", 1)
    add_table(
        doc,
        ["Say this", "Avoid this"],
        [
            ["The descriptive genomic results are partly corroborated.", "The thesis is fake."],
            ["The headline transition and causal claims are not yet reproducible from available materials.", "The results are illegitimate."],
            ["The differences may be due to denominator and clinical-definition choices.", "She must have manipulated the data."],
            ["Grey-zone handling needs explicit reconciliation.", "The model is definitely wrong."],
            ["Antibiotic claims should be treated as associations unless causal design is supplied.", "Antibiotics definitely do not matter."],
        ],
        [4680, 4680],
    )

    add_heading(doc, "Questions to ask", 1)
    questions = [
        "Can we see the row-level wgMLST isolate and pair-distance table, including comparable loci and missing loci?",
        "Which exact E. coli wgMLST scheme version was used, and why does the thesis report 4,512 loci when the cited benchmark reports 4,503?",
        "Were the 26 grey-zone pairs excluded, treated as uncertain, or counted as replacement in the trajectory and Cox analyses?",
        "Why are there 368 consecutive wgMLST pairs but 394 WGS/VF-linked transitions in my current pipeline?",
        "Why does the thesis report 34 ASB-to-UTI pairs but also use 26 ASB-to-UTI pairs in another transition partition?",
        "How were UTI episodes defined: event label, Castor UTI, culture/symptom rule, or sequenced UTI isolate?",
        "How were repeated pairs from the same resident handled statistically?",
        "For the Cox model, what was the time origin, event definition, censoring rule, clustering method, proportional-hazards diagnostic and missing-exposure handling?",
        "Can we see the Castor antibiotic extract with indication, drug, start/end dates and linkage rules?",
        "Can causal language be softened unless a causal model/design is provided?",
    ]
    for question in questions:
        add_numbered(doc, question)

    add_heading(doc, "Evidence still needed", 1)
    for item in [
        "wgMLST pair-distance table and isolate-level allele profiles.",
        "Assembly/QC manifest with assembler choice, coverage, polishing and QC provenance.",
        "Exact SeqSphere scheme name/version/date and target export.",
        "Row-level ASB/UTI transition and recurrence tables.",
        "Castor antibiotic extract and model-ready antibiotic variables.",
        "Model formulas, code, package/session information and missing supplementary tables.",
    ]:
        add_bullet(doc, item)

    doc.save(DOCX_OUT)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    MD_OUT.write_text(build_markdown(), encoding="utf-8")
    build_docx()
    print(f"Wrote {DOCX_OUT}")
    print(f"Wrote {MD_OUT}")


if __name__ == "__main__":
    main()
